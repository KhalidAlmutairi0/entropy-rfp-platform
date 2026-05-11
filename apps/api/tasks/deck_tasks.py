"""Celery tasks for proposal deck generation (Word template fill + LibreOffice PDF conversion).

Flow:
  1. Load master .docx template from MinIO (or local path configured via DECK_TEMPLATE_PATH).
  2. Use OpenAI gpt-4o (or strongest available) to generate content for each variable placeholder,
     grounded in the ingested RFP text extracted at analysis time.
  3. Fill placeholders in the .docx using python-docx.
  4. Convert the filled .docx to PDF via LibreOffice --headless.
  5. Store the PDF in MinIO under decks/{rfp_id}/proposal.pdf.
  6. Update rfp.deck_status = "READY" and rfp.deck_pdf_path.

NOTE: The placeholder schema (which sections are variable vs. static) is derived from
the example .docx pairs provided by the user. Until those examples are provided,
the PLACEHOLDER_SCHEMA below acts as the configuration point — update it once the
examples are analysed.
"""

import asyncio
import io
import json
import os
import subprocess
import tempfile
from pathlib import Path

import structlog

from core.celery_app import celery_app

logger = structlog.get_logger()

# ── Placeholder schema ─────────────────────────────────────────────────────────
# Each key is the exact placeholder text inside the .docx (e.g. "{{EXECUTIVE_SUMMARY}}").
# The value is an instruction string sent to the model for that section.
# UPDATE THIS when the user provides before/after .docx examples.
PLACEHOLDER_SCHEMA: dict[str, str] = {
    "{{EXECUTIVE_SUMMARY}}": (
        "Write a concise Arabic executive summary (150–200 words) tailored to this specific RFP. "
        "Highlight Entropy's fit, key differentiators, and expected outcomes. "
        "Ground every claim in the RFP text. Do not hallucinate."
    ),
    "{{UNDERSTANDING_OF_REQUIREMENTS}}": (
        "Demonstrate deep understanding of the client's stated requirements. "
        "Reference specific sections or page numbers from the RFP. "
        "Write in Arabic, 150–200 words."
    ),
    "{{PROPOSED_METHODOLOGY}}": (
        "Describe Entropy's proposed methodology for this engagement. "
        "Align with the RFP's technical scope. Arabic, 150–200 words."
    ),
    "{{PROJECT_TIMELINE}}": (
        "Provide a high-level milestone plan (phases, durations) appropriate for this RFP. "
        "Arabic, bullet format, max 8 milestones."
    ),
    "{{PRICING_SUMMARY}}": (
        "Leave as [يُحدَّد لاحقاً — To be determined]. Do NOT generate pricing — this section requires manual completion."
    ),
}

# Path to master .docx template. Override via env var DECK_TEMPLATE_PATH.
_TEMPLATE_ENV = os.getenv("DECK_TEMPLATE_PATH", "")


@celery_app.task(bind=True, name="tasks.deck_tasks.generate_deck_task")
def generate_deck_task(self, rfp_id: str) -> dict:
    return asyncio.run(_generate_deck_async(rfp_id))


async def _generate_deck_async(rfp_id: str) -> dict:
    from core.database import AsyncSessionLocal
    from models.rfp import RFP
    from services.storage import StorageService
    from sqlalchemy import select
    from sqlalchemy.orm import selectinload

    storage = StorageService.get_instance()

    async with AsyncSessionLocal() as db:
        result = await db.execute(
            select(RFP).options(selectinload(RFP.files)).where(RFP.id == rfp_id)
        )
        rfp = result.scalar_one_or_none()
        if not rfp:
            logger.error("Deck task: RFP not found", rfp_id=rfp_id)
            return {"error": "RFP not found"}

        rfp.deck_status = "GENERATING"
        await db.commit()

        try:
            # 1. Load template
            template_bytes = await _load_template(storage)
            if template_bytes is None:
                rfp.deck_status = "FAILED"
                await db.commit()
                return {"error": "Master template not found. Set DECK_TEMPLATE_PATH env var or upload to MinIO at templates/master_proposal.docx"}

            # 2. Extract RFP text for grounding
            rfp_text = await _extract_rfp_text(rfp, storage)

            # 3. Generate content for each placeholder via OpenAI
            filled_values = await _generate_placeholder_content(rfp, rfp_text)

            # 4. Fill the .docx template
            filled_docx_bytes = _fill_docx_template(template_bytes, filled_values)

            # 5. Convert to PDF via LibreOffice --headless
            pdf_bytes = _convert_docx_to_pdf(filled_docx_bytes)

            # 6. Store PDF in MinIO
            pdf_path = f"decks/{rfp_id}/proposal.pdf"
            await storage.upload_file(pdf_bytes, pdf_path, "application/pdf")

            rfp.deck_pdf_path = pdf_path
            rfp.deck_status = "READY"
            await db.commit()

            logger.info("Deck generated successfully", rfp_id=rfp_id, path=pdf_path)
            return {"rfp_id": rfp_id, "pdf_path": pdf_path, "status": "READY"}

        except Exception as exc:
            logger.error("Deck generation failed", rfp_id=rfp_id, error=str(exc))
            rfp.deck_status = "FAILED"
            await db.commit()
            raise


async def _load_template(storage) -> bytes | None:
    """Load the master .docx template from local path or MinIO."""
    # 1. Local filesystem override (fastest for dev)
    if _TEMPLATE_ENV and Path(_TEMPLATE_ENV).exists():
        return Path(_TEMPLATE_ENV).read_bytes()

    # 2. MinIO
    try:
        return await storage.download_file("templates/master_proposal.docx")
    except Exception:
        pass

    # 3. Bundled fallback path (next to this file)
    bundled = Path(__file__).parent / "master_proposal.docx"
    if bundled.exists():
        return bundled.read_bytes()

    return None


async def _extract_rfp_text(rfp, storage) -> str:
    """Return the ingested RFP text for use as grounding context."""
    from tasks.ingestion_tasks import _extract_text

    blocks: list[str] = []
    for rfp_file in rfp.files:
        try:
            content = await storage.download_file(rfp_file.storage_path)
            text = await _extract_text(content, rfp_file.mime_type or "", rfp_file.filename)
            if text.strip():
                blocks.append(f"=== {rfp_file.filename} ===\n{text.strip()[:10000]}")
        except Exception as exc:
            logger.warning("Failed to extract text for deck", file=rfp_file.filename, error=str(exc))

    return "\n\n".join(blocks)[:30000]  # Cap at 30k chars for context window


async def _generate_placeholder_content(rfp, rfp_text: str) -> dict[str, str]:
    """Call OpenAI to generate content for each variable placeholder."""
    import openai

    client = openai.AsyncOpenAI()

    # Select strongest available model
    model = _select_model(client)

    system_prompt = (
        "You are an expert Arabic proposal writer for Entropy, a Saudi AI company. "
        "You generate precise, professional proposal content grounded in the provided RFP text. "
        "Never hallucinate facts. Cite the RFP where relevant. "
        "Output ONLY the requested content, no preamble."
    )

    rfp_header = (
        f"RFP Title: {rfp.title_ar or rfp.title_en or 'N/A'}\n"
        f"Agency: {rfp.agency or 'N/A'}\n"
        f"Tender: {rfp.tender_number or 'N/A'}\n"
    )

    filled: dict[str, str] = {}

    for placeholder, instruction in PLACEHOLDER_SCHEMA.items():
        # Pricing placeholder is never generated
        if "يُحدَّد لاحقاً" in instruction or "To be determined" in instruction:
            filled[placeholder] = "[يُحدَّد لاحقاً — To be determined]"
            continue

        user_prompt = (
            f"{rfp_header}\n\n"
            f"RFP Content (excerpt):\n{rfp_text[:8000]}\n\n"
            f"Task: {instruction}"
        )

        try:
            response = await client.chat.completions.create(
                model=model,
                temperature=0.3,  # Max 0.3 for prose as per spec
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                max_tokens=600,
            )
            content = response.choices[0].message.content or ""

            # Second review pass: verify no hallucinations
            content = await _review_generated_content(client, model, placeholder, content, rfp_text)
            filled[placeholder] = content.strip()

        except Exception as exc:
            logger.error("OpenAI call failed for placeholder", placeholder=placeholder, error=str(exc))
            filled[placeholder] = f"[فشل التوليد — {placeholder}]"

    return filled


async def _review_generated_content(client, model: str, placeholder: str, content: str, rfp_text: str) -> str:
    """Second-pass review: check for hallucinations or missing grounding."""
    review_prompt = (
        f"Review the following generated proposal section for placeholder '{placeholder}'.\n\n"
        f"Generated content:\n{content}\n\n"
        f"RFP source (excerpt):\n{rfp_text[:4000]}\n\n"
        "Check: (1) Are all factual claims grounded in the RFP? (2) Any hallucinated numbers, names, or dates?\n"
        "If the content is sound, return it unchanged.\n"
        "If there are issues, return a corrected version that removes any ungrounded claims.\n"
        "Return ONLY the final content, no explanation."
    )

    try:
        response = await client.chat.completions.create(
            model=model,
            temperature=0,
            messages=[{"role": "user", "content": review_prompt}],
            max_tokens=600,
        )
        return response.choices[0].message.content or content
    except Exception:
        return content  # Fall back to original if review fails


def _select_model(client) -> str:
    """Return the strongest available OpenAI model at runtime.
    Preference order: o3 > o1 > gpt-4.5 > gpt-4o. Never use gpt-3.5 or non-flagship models.
    """
    preferred = ["o3", "o1", "gpt-4.5", "gpt-4o"]
    # Try to list available models; fall back to gpt-4o if API call fails
    try:
        import openai
        available = {m.id for m in openai.OpenAI().models.list().data}
        for model in preferred:
            if model in available:
                return model
    except Exception:
        pass
    return "gpt-4o"


def _fill_docx_template(template_bytes: bytes, filled_values: dict[str, str]) -> bytes:
    """Replace placeholder strings in the .docx using python-docx.
    Preserves all static content, fonts, tables, images, and RTL direction exactly.
    """
    from docx import Document

    doc = Document(io.BytesIO(template_bytes))

    def replace_in_paragraphs(paragraphs):
        for para in paragraphs:
            for placeholder, value in filled_values.items():
                if placeholder in para.text:
                    # Replace across runs while preserving run formatting
                    full_text = "".join(r.text for r in para.runs)
                    if placeholder in full_text:
                        new_text = full_text.replace(placeholder, value)
                        # Put all text in first run, clear the rest
                        if para.runs:
                            para.runs[0].text = new_text
                            for run in para.runs[1:]:
                                run.text = ""

    # Replace in body paragraphs
    replace_in_paragraphs(doc.paragraphs)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs)

    # Replace in headers and footers
    for section in doc.sections:
        replace_in_paragraphs(section.header.paragraphs)
        replace_in_paragraphs(section.footer.paragraphs)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    """Convert a .docx file to PDF using LibreOffice --headless.
    Preserves Arabic RTL text direction, images, tables, and styled text.
    Requires LibreOffice to be installed on the server.
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        input_path = Path(tmpdir) / "input.docx"
        input_path.write_bytes(docx_bytes)

        result = subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", tmpdir,
                str(input_path),
            ],
            capture_output=True,
            text=True,
            timeout=120,
        )

        if result.returncode != 0:
            raise RuntimeError(
                f"LibreOffice conversion failed (exit {result.returncode}): {result.stderr}"
            )

        output_path = Path(tmpdir) / "input.pdf"
        if not output_path.exists():
            raise RuntimeError("LibreOffice did not produce an output PDF")

        return output_path.read_bytes()
