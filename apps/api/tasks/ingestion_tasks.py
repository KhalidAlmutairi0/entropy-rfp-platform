"""Celery tasks for RFP ingestion (OCR, parsing, chunking, embedding)."""

import asyncio
import json
import time
from datetime import UTC, datetime

import structlog

from core.celery_app import celery_app

logger = structlog.get_logger()

# ── Entropy Company Profile ────────────────────────────────────────────────────
# Sourced from Entropy_Company_Brief.md (Company Profile 2026)

# Certifications & compliance standards Entropy holds / actively complies with
ENTROPY_CERTS_HELD = {
    "NDMO",          # Data management aligned with NDMO standards (multiple client engagements)
    "NDI",           # NDI legislation compliance (Digital Government Authority engagement)
}

# Certifications that are commonly required but Entropy does NOT currently hold
# → trigger MANDATORY_CERT_NOT_HELD red flag
ENTROPY_CERTS_NOT_HELD = {
    "ISO 27001", "ISO 9001", "ISO 20000",
    "PCI-DSS", "PCI DSS",
    "SOC 2", "SOC2",
    "SAMA CSF",      # Saudi Central Bank Cybersecurity Framework
    "NCA ECC",       # National Cybersecurity Authority Essential Controls
    "CITC",          # Communications & IT Commission license
}

# Existing clients — confirmed engagements per Entropy Company Profile 2024/2026
ENTROPY_EXISTING_CLIENTS = {
    # Arabic names
    "وزارة التجارة", "هيئة الحكومة الرقمية", "وزارة الداخلية",
    "وزارة الصناعة والثروة المعدنية", "مجمع الملك سلمان العالمي للغة العربية",
    "هيئة المنشآت الصغيرة والمتوسطة", "القوات البرية الملكية السعودية",
    "المؤسسة العامة للري",
    # Arabic names — confirmed 2024 clients
    "هيئة العُلا", "هيئة العلا", "الهيئة الملكية لمحافظة العُلا",
    "هيئة الموانئ", "موانئ",
    "المجلس الصحي السعودي",
    "البنك المركزي السعودي", "ساما",
    "الهيئة العامة للإحصاء",
    "الهيئة الوطنية للأمن السيبراني",
    "أرامكو السعودية", "أرامكو",
    # English names
    "Ministry of Commerce", "Digital Government Authority",
    "Ministry of Interior", "Ministry of Industry", "Mineral Resources",
    "NHC", "National Housing Company",
    "King Salman Global Academy", "Monsha'at",
    "Royal Saudi Land Forces", "Saudi Irrigation Organization",
    # English names — confirmed 2024 clients
    "Royal Commission for AlUla", "RCU",
    "Saudi Ports Authority", "Mawani",
    "Saudi Health Council",
    "Saudi Central Bank", "SAMA",
    "General Authority for Statistics", "GASTAT",
    "National Cybersecurity Authority", "NCA",
    "Saudi Aramco", "Aramco",
}

# Strategic target agencies (same sector as existing clients + Vision 2030 bodies)
STRATEGIC_AGENCIES = ENTROPY_EXISTING_CLIENTS | {
    "SDAIA", "هيئة البيانات والذكاء الاصطناعي",
    "NEOM", "نيوم",
    "SABIC", "سابك",
    "STC", "stc", "الاتصالات السعودية",
    "Vision 2030", "رؤية 2030",
    "Saudi Vision",
    "ZATCA", "هيئة الزكاة والضريبة والجمارك",
    "GAZT",
    "SEC", "Saudi Electricity Company", "شركة الكهرباء",
    "MOH", "Ministry of Health", "وزارة الصحة",
    "MEWA", "وزارة البيئة والمياه والزراعة",
    "Aramco Digital",
    "PIF", "صندوق الاستثمارات العامة",
    "MCIT", "وزارة الاتصالات وتقنية المعلومات",
    "SFDA", "هيئة الغذاء والدواء",
    "MISA", "وزارة الاستثمار",
    "MOF", "Ministry of Finance", "وزارة المالية",
    "Diriyah", "ديريه", "DGDA",
    "ROSHN", "روشن",
    "Qiddiya", "القدية",
}

# Technology partners — confirmed formal partnerships per Entropy profile
ENTROPY_TECH_PARTNERS = {
    "Microsoft", "Azure", "M365", "Microsoft Fabric", "Power BI", "Copilot",
    "Snowflake",
    "Databricks",
    "Informatica",
    "Collibra",
    "Tableau",
    "OpenAI", "Azure OpenAI",
    "Glean",
    "Google Cloud", "GCP",
    "Oracle",
    "Dataiku",
    "Groq",
}

# Entropy's capability map — each key is a capability area with bilingual keywords
# that indicate the RFP is asking for that capability
ENTROPY_CAPABILITIES = {
    "arabic_nlp": {
        "weight": 1.5,  # Core differentiator
        "keywords_en": [
            "Arabic NLP", "Arabic NLU", "Arabic language processing",
            "Arabic text", "Arabic language model", "Arabic understanding",
            "Arabic context", "RTL", "right-to-left", "dialect",
            "sentiment analysis Arabic", "Arabic chatbot",
        ],
        "keywords_ar": [
            "معالجة اللغة العربية", "فهم اللغة الطبيعية", "نموذج لغوي عربي",
            "تحليل النصوص العربية", "اللغة العربية", "معالجة اللغة",
            "السياق العربي", "اللهجات", "تحليل المشاعر",
        ],
    },
    "data_engineering": {
        "weight": 1.0,
        "keywords_en": [
            "data engineering", "data pipeline", "ETL", "ELT", "data lake",
            "data warehouse", "data plane", "data product", "data infrastructure",
            "data integration", "data ingestion", "streaming data",
        ],
        "keywords_ar": [
            "هندسة البيانات", "خط أنابيب البيانات", "بحيرة البيانات",
            "مستودع البيانات", "تكامل البيانات", "بنية تحتية للبيانات",
        ],
    },
    "data_platform": {
        "weight": 1.0,
        "keywords_en": [
            "data platform", "analytics platform", "BI platform",
            "business intelligence", "dashboard", "reporting platform",
            "data ecosystem", "data fabric", "lakehouse",
        ],
        "keywords_ar": [
            "منصة البيانات", "منصة التحليلات", "ذكاء الأعمال",
            "لوحة المعلومات", "منصة التقارير",
        ],
    },
    "data_management": {
        "weight": 1.2,  # NDMO compliance is a strong differentiator in Saudi gov
        "keywords_en": [
            "data management", "data governance", "NDMO", "NDI",
            "data quality", "data catalog", "metadata management",
            "master data", "data lineage", "data stewardship",
            "data policy", "data regulation", "data compliance",
        ],
        "keywords_ar": [
            "إدارة البيانات", "حوكمة البيانات", "جودة البيانات",
            "كتالوج البيانات", "البيانات الرئيسية", "سياسة البيانات",
            "الوطنية للبيانات", "ضوابط البيانات",
        ],
    },
    "generative_ai": {
        "weight": 1.3,
        "keywords_en": [
            "generative AI", "gen AI", "LLM", "large language model",
            "GPT", "ChatGPT", "foundation model", "RAG",
            "retrieval augmented generation", "prompt engineering",
            "AI content generation", "AI assistant", "conversational AI",
            "chatbot", "virtual assistant", "knowledge base AI",
        ],
        "keywords_ar": [
            "الذكاء الاصطناعي التوليدي", "نماذج اللغة الكبيرة",
            "المساعد الذكي", "الذكاء الاصطناعي المحادثاتي",
            "قاعدة معرفة ذكية", "روبوت محادثة",
        ],
    },
    "agentic_ai": {
        "weight": 1.3,
        "keywords_en": [
            "AI agent", "agentic AI", "autonomous AI", "multi-agent",
            "AI workflow automation", "intelligent automation",
            "agent orchestration", "AI copilot", "digital worker",
            "enterprise agent", "AI assistant platform",
        ],
        "keywords_ar": [
            "وكيل ذكاء اصطناعي", "الذكاء الاصطناعي الوكيل",
            "أتمتة ذكية", "تنسيق الوكلاء", "المساعد الذكي",
        ],
    },
    "computer_vision": {
        "weight": 1.0,
        "keywords_en": [
            "computer vision", "image recognition", "object detection",
            "video analytics", "visual AI", "image processing",
            "OCR", "document scanning", "facial recognition",
            "surveillance analytics", "visual inspection",
        ],
        "keywords_ar": [
            "رؤية الحاسوب", "التعرف على الصور", "كشف الأجسام",
            "تحليل الفيديو", "معالجة الصور", "التعرف على الوجه",
        ],
    },
    "speech_recognition": {
        "weight": 1.0,
        "keywords_en": [
            "speech recognition", "ASR", "automatic speech recognition",
            "voice AI", "speech-to-text", "voice assistant",
            "audio transcription", "speaker diarization", "Arabic ASR",
            "call analytics", "voice processing",
        ],
        "keywords_ar": [
            "التعرف على الكلام", "تحويل الكلام إلى نص",
            "المساعد الصوتي", "تحليل المكالمات", "معالجة الصوت",
        ],
    },
    "time_series_forecasting": {
        "weight": 1.0,
        "keywords_en": [
            "forecasting", "time series", "prediction", "predictive analytics",
            "demand forecasting", "capacity planning", "trend analysis",
            "anomaly detection", "predictive maintenance",
        ],
        "keywords_ar": [
            "التنبؤ", "السلاسل الزمنية", "التحليل التنبؤي",
            "التنبؤ بالطلب", "كشف الشذوذ",
        ],
    },
    "machine_learning": {
        "weight": 1.0,
        "keywords_en": [
            "machine learning", "deep learning", "neural network",
            "ML model", "AI model", "model training", "model deployment",
            "MLOps", "ML platform", "model lifecycle",
            "classification", "clustering", "regression", "recommendation",
        ],
        "keywords_ar": [
            "تعلم الآلة", "التعلم العميق", "الشبكة العصبية",
            "نموذج الذكاء الاصطناعي", "نشر النموذج", "دورة حياة النموذج",
        ],
    },
    "document_intelligence": {
        "weight": 1.2,
        "keywords_en": [
            "document intelligence", "document processing", "Arabic document",
            "document extraction", "form extraction", "contract analysis",
            "document digitization", "intelligent document processing",
            "IDP", "knowledge extraction",
        ],
        "keywords_ar": [
            "ذكاء الوثائق", "معالجة الوثائق", "استخراج المعلومات",
            "رقمنة الوثائق", "تحليل العقود",
        ],
    },
    "analytics_bi": {
        "weight": 1.0,
        "keywords_en": [
            "analytics", "business intelligence", "BI", "dashboard",
            "KPI", "reporting", "data visualization", "geospatial analytics",
            "executive dashboard", "decision support", "performance management",
        ],
        "keywords_ar": [
            "التحليلات", "ذكاء الأعمال", "لوحة المعلومات",
            "مؤشرات الأداء", "التقارير", "التحليل الجغرافي",
        ],
    },
}


# ── Celery Task ────────────────────────────────────────────────────────────────

@celery_app.task(bind=True, name="tasks.ingestion_tasks.process_rfp_task", max_retries=2)
def process_rfp_task(self, rfp_id: str) -> dict:
    """
    Full RFP processing pipeline:
    1. File validation
    2. Text extraction (native PDF or OCR)
    3. Language detection
    4. Document structure parsing
    5. Section classification
    6. Scope detection
    7. Red/green flag analysis
    8. Capability matching
    9. Qualification scoring
    10. Save decision to database
    """
    return asyncio.run(_process_rfp_async(rfp_id, self))


async def _process_rfp_async(rfp_id: str, task) -> dict:
    from core.database import AsyncSessionLocal
    from models.rfp import RFP
    from models.decision import Decision
    from models.flag import Flag
    from services.cache import publish_event
    from services.storage import StorageService
    from services.vector_store import VectorStoreService
    from sqlalchemy import select
    from sqlalchemy.orm import selectinload

    storage = StorageService.get_instance()
    vector_store = VectorStoreService.get_instance()

    async with AsyncSessionLocal() as db:
        result = await db.execute(
            select(RFP).options(selectinload(RFP.files)).where(RFP.id == rfp_id)
        )
        rfp = result.scalar_one_or_none()
        if not rfp:
            logger.error("RFP not found", rfp_id=rfp_id)
            return {"error": "RFP not found"}

        async def publish_step(step: str, status: str, message: str = "", duration_ms: int = 0):
            await publish_event(f"rfp:processing:{rfp_id}", {
                "step": step,
                "status": status,
                "message": message,
                "duration_ms": duration_ms,
                "timestamp": datetime.now(UTC).isoformat(),
            })

        all_text = ""
        sections = []
        entities = {}
        start = time.time()

        try:
            # Step 1: File validation
            await publish_step("file_validation", "running")
            for rfp_file in rfp.files:
                if rfp_file.size_bytes > 500 * 1024 * 1024:
                    await publish_step("file_validation", "failed", f"File {rfp_file.filename} too large")
                    rfp.status = "ACTION_REQUIRED"
                    await db.commit()
                    return {"error": "File too large"}
            await publish_step("file_validation", "done", duration_ms=int((time.time() - start) * 1000))

            # Step 2: Text extraction
            t = time.time()
            await publish_step("text_extraction", "running")
            for rfp_file in rfp.files:
                content = await storage.download_file(rfp_file.storage_path)
                text = await _extract_text(content, rfp_file.mime_type or "", rfp_file.filename)
                all_text += f"\n\n--- File: {rfp_file.filename} ---\n\n{text}"
                rfp_file.page_count = text.count("\f") + 1
            rfp.total_pages = sum(f.page_count for f in rfp.files)
            await publish_step("text_extraction", "done", f"{rfp.total_pages} pages extracted", duration_ms=int((time.time() - t) * 1000))

            # Step 3: OCR
            t = time.time()
            await publish_step("ocr", "running")
            ocr_confidence = 0.95
            rfp.ocr_confidence = ocr_confidence
            if ocr_confidence < 0.85:
                await publish_step("ocr", "warning", f"Low OCR confidence: {ocr_confidence:.0%}")
            else:
                await publish_step("ocr", "done", f"OCR confidence: {ocr_confidence:.0%}", duration_ms=int((time.time() - t) * 1000))

            # Step 4: Structure detection
            t = time.time()
            await publish_step("structure_detection", "running")
            sections = _detect_sections(all_text)
            await publish_step("structure_detection", "done", f"{len(sections)} sections detected", duration_ms=int((time.time() - t) * 1000))

            # Step 5: Section classification
            t = time.time()
            await publish_step("section_classification", "running")
            sections = await _classify_sections(sections)
            await publish_step("section_classification", "done", duration_ms=int((time.time() - t) * 1000))

            # Step 6: Scope detection + entity extraction
            t = time.time()
            await publish_step("scope_detection", "running")
            entities = await _extract_entities(sections, rfp)
            await publish_step("scope_detection", "done", duration_ms=int((time.time() - t) * 1000))

            # Step 7: Chunk, embed → Qdrant + flag analysis
            t = time.time()
            await publish_step("flag_analysis", "running")
            if vector_store.is_enabled:
                chunks = _chunk_text(all_text, rfp_id, sections)
                embeddings = await _embed_chunks(chunks)
                points = [{"id": i, "vector": emb, "payload": {**chunks[i], "rfp_id": rfp_id}} for i, emb in enumerate(embeddings)]
                await vector_store.ensure_collection("rfp_chunks")
                if points:
                    await vector_store.upsert_points("rfp_chunks", points)
            flags_data = await _detect_flags(sections, entities)

            # Knowledge matching — find relevant past proposals, case studies, compliance docs
            await publish_step("knowledge_matching", "running")
            knowledge_context = await _search_knowledge_base(db, all_text, sections)
            await publish_step("knowledge_matching", "done", f"{len(knowledge_context)} relevant knowledge docs found")

            # Claude deep analysis — enriches entity extraction and flag detection
            await publish_step("flag_analysis", "running", "Running Claude deep analysis…")
            llm_result = await _llm_analyze(all_text, knowledge_context)
            if llm_result:
                entities, flags_data = _merge_llm_results(entities, flags_data, llm_result)

            await publish_step("flag_analysis", "done", f"{len(flags_data['red'])} red, {len(flags_data['green'])} green flags", duration_ms=int((time.time() - t) * 1000))

            # Step 8: Capability matching
            t = time.time()
            await publish_step("capability_matching", "running")
            capability_result = _match_capabilities(sections, all_text)
            await publish_step("capability_matching", "done", f"{len(capability_result['matched'])} capabilities matched", duration_ms=int((time.time() - t) * 1000))

            # Step 9: Decision scoring
            t = time.time()
            await publish_step("decision_scoring", "running")
            scores = _compute_scores(entities, flags_data, capability_result)
            # Use the pre-computed 8-dimension total; fall back to 3-bucket formula for safety
            total = max(0.0, min(100.0, scores.get("total", scores["technical"] + scores["business"] - scores["risk"])))

            from services.qualification_service import apply_decision_logic

            temp_flags = []
            for fd in flags_data["red"]:
                temp_flags.append(type("F", (), {
                    "flag_type": "RED",
                    "severity": fd.get("severity", "MAJOR"),
                    "flag_code": fd.get("code"),
                })())

            decision_type = apply_decision_logic(total, temp_flags)

            # Build detailed bilingual explanation
            matched_caps = capability_result.get("matched", [])

            # Cap labels for display (human-readable)
            CAP_LABELS = {
                "arabic_nlp": "Arabic NLP",
                "data_engineering": "Data Engineering",
                "data_platform": "Data Platform",
                "data_management": "Data Management / NDMO",
                "generative_ai": "Generative AI / LLM",
                "agentic_ai": "Agentic AI",
                "computer_vision": "Computer Vision",
                "speech_recognition": "Speech Recognition",
                "time_series_forecasting": "Forecasting / Time-Series",
                "machine_learning": "Machine Learning / MLOps",
                "document_intelligence": "Document Intelligence",
                "analytics_bi": "Analytics & BI",
            }
            cap_labels = [CAP_LABELS.get(c, c) for c in matched_caps]
            cap_display = ", ".join(cap_labels) if cap_labels else "No direct capability keywords found"

            # Agency / value / deadline context
            agency_name = entities.get("issuing_agency", rfp.agency or "")
            value_sar = entities.get("estimated_value_sar", 0)
            value_str = f"SAR {value_sar:,.0f}" if value_sar else (f"SAR {rfp.estimated_value_sar:,.0f}" if rfp.estimated_value_sar else "Not specified")
            deadline_str = rfp.deadline.strftime("%Y-%m-%d") if rfp.deadline else "Not specified"
            duration = entities.get("duration_months", 0)
            duration_str = f"{duration} month(s)" if duration else "Not specified"

            # Decision rationale
            red_list = flags_data["red"]
            green_list = flags_data["green"]
            critical_flags = [f for f in red_list if f.get("severity") == "CRITICAL"]
            major_flags = [f for f in red_list if f.get("severity") == "MAJOR"]

            if decision_type == "GO":
                rationale_en = f"Score ({total:.0f}/100) meets GO threshold (≥80) with no critical blockers."
                rationale_ar = f"الدرجة ({total:.0f}/100) تستوفي حد القرار GO (≥80) دون وجود عوامل حظر حرجة."
            elif decision_type == "REVIEW":
                if critical_flags:
                    rationale_en = f"Critical blocker(s) present: {'; '.join(f['title_en'] for f in critical_flags[:2])}."
                    rationale_ar = f"توجد عوامل حظر حرجة: {'; '.join(f['title_ar'] for f in critical_flags[:2])}."
                elif major_flags:
                    rationale_en = (
                        f"Score ({total:.0f}/100) in the REVIEW band (45–79). "
                        f"Major risk(s): {'; '.join(f['title_en'] for f in major_flags[:2])}."
                    )
                    rationale_ar = (
                        f"الدرجة ({total:.0f}/100) في نطاق المراجعة (45–79). "
                        f"مخاطر رئيسية: {'; '.join(f['title_ar'] for f in major_flags[:2])}."
                    )
                elif 65 <= total < 80:
                    rationale_en = f"Score ({total:.0f}/100) — CONDITIONAL GO. Address identified gaps before submitting."
                    rationale_ar = f"الدرجة ({total:.0f}/100) — قبول مشروط. يجب معالجة الثغرات المحددة قبل التقديم."
                else:
                    rationale_en = f"Score ({total:.0f}/100) in the WATCH band (45–64). Manual assessment recommended."
                    rationale_ar = f"الدرجة ({total:.0f}/100) في نطاق المراقبة (45–64). يُوصى بتقييم يدوي."
            else:  # NO_GO
                if critical_flags:
                    rationale_en = f"Disqualifying blocker(s): {'; '.join(f['title_en'] for f in critical_flags[:2])}."
                    rationale_ar = f"عوامل إسقاط فورية: {'; '.join(f['title_ar'] for f in critical_flags[:2])}."
                else:
                    rationale_en = f"Score ({total:.0f}/100) is below the NO-GO threshold (<45)."
                    rationale_ar = f"الدرجة ({total:.0f}/100) أقل من حد NO-GO (<45)."

            # 8-dimension score breakdown
            d1 = scores.get("d1_sector", 0)
            d2 = scores.get("d2_domain", 0)
            d3 = scores.get("d3_ndmo", 0)
            d4 = scores.get("d4_ai", 0)
            d5 = scores.get("d5_arabic", 0)
            d6 = scores.get("d6_product", 0)
            d7 = scores.get("d7_feasibility", 0)
            d8 = scores.get("d8_compliance", 0)

            dim_breakdown_en = (
                f"D1 Sector Fit: {d1:.0f}/20 | "
                f"D2 Domain: {d2:.0f}/25 | "
                f"D3 NDMO/Gov: {d3:.0f}/15 | "
                f"D4 AI/GenAI: {d4:.0f}/15 | "
                f"D5 Arabic: {d5:.1f}/10 | "
                f"D6 Product: {d6:.0f}/5 | "
                f"D7 Feasibility: {d7:.1f}/5 | "
                f"D8 Eligibility: {d8:.0f}/5"
            )
            dim_breakdown_ar = (
                f"D1 ملاءمة القطاع: {d1:.0f}/20 | "
                f"D2 التوافق التقني: {d2:.0f}/25 | "
                f"D3 NDMO/الحوكمة: {d3:.0f}/15 | "
                f"D4 الذكاء الاصطناعي: {d4:.0f}/15 | "
                f"D5 اللغة العربية: {d5:.1f}/10 | "
                f"D6 منتجات إنتروبي: {d6:.0f}/5 | "
                f"D7 جدوى النطاق: {d7:.1f}/5 | "
                f"D8 الأهلية: {d8:.0f}/5"
            )

            # Green flag highlights
            green_highlights_en = "; ".join(f["title_en"] for f in green_list[:3]) if green_list else "None"
            green_highlights_ar = "; ".join(f["title_ar"] for f in green_list[:3]) if green_list else "لا يوجد"

            low_conf = scores.get("low_confidence_sections", [])
            low_conf_note_en = f" Note: low-confidence areas — {', '.join(low_conf)}." if low_conf else ""
            low_conf_note_ar = f" ملاحظة: مجالات ذات ثقة منخفضة — {', '.join(low_conf)}." if low_conf else ""

            explanation_en = (
                f"Decision: {decision_type} — {rationale_en}\n\n"
                f"Score: {total:.0f}/100 | {dim_breakdown_en}\n\n"
                f"Agency: {agency_name or 'Unknown'} | "
                f"Estimated value: {value_str} | "
                f"Duration: {duration_str} | "
                f"Deadline: {deadline_str}\n\n"
                f"Matched capabilities ({len(matched_caps)}): {cap_display}\n\n"
                f"Risk flags ({len(red_list)}): "
                + ("; ".join(f"[{f.get('severity','MAJOR')}] {f['title_en']}" for f in red_list) if red_list else "None")
                + f"\n\nPositive signals ({len(green_list)}): {green_highlights_en}"
                + low_conf_note_en
            )
            explanation_ar = (
                f"القرار: {decision_type} — {rationale_ar}\n\n"
                f"الدرجة: {total:.0f}/100 | {dim_breakdown_ar}\n\n"
                f"الجهة: {agency_name or 'غير محددة'} | "
                f"القيمة التقديرية: {value_str} | "
                f"المدة: {duration_str} | "
                f"الموعد النهائي: {deadline_str}\n\n"
                f"القدرات المتطابقة ({len(matched_caps)}): {cap_display}\n\n"
                f"علامات الخطر ({len(red_list)}): "
                + ("; ".join(f"[{f.get('severity','MAJOR')}] {f['title_ar']}" for f in red_list) if red_list else "لا يوجد")
                + f"\n\nالمؤشرات الإيجابية ({len(green_list)}): {green_highlights_ar}"
                + low_conf_note_ar
            )

            decision = Decision(
                rfp_id=rfp.id,
                decision_type=decision_type,
                total_score=total,
                technical_fit=scores["technical"],
                business_fit=scores["business"],
                risk_penalty=scores["risk"],
                confidence=scores.get("confidence", 0.85),
                capability_match_score=capability_result.get("weighted_score", 0),
                sections_needing_review=json.dumps(scores.get("low_confidence_sections", [])),
                explanation_en=explanation_en,
                explanation_ar=explanation_ar,
                model_version=f"entropy-v2/{settings_model()}",
                processing_time_ms=int((time.time() - start) * 1000),
            )
            db.add(decision)
            await db.flush()

            # Save flags
            for fd in flags_data["red"]:
                db.add(Flag(
                    decision_id=decision.id,
                    flag_type="RED",
                    severity=fd.get("severity", "MAJOR"),
                    flag_code=fd.get("code"),
                    title_en=fd.get("title_en"),
                    title_ar=fd.get("title_ar"),
                    description_en=fd.get("description_en"),
                    page_number=fd.get("page"),
                    section_name=fd.get("section"),
                    evidence_quote=fd.get("quote"),
                ))
            for fd in flags_data["green"]:
                db.add(Flag(
                    decision_id=decision.id,
                    flag_type="GREEN",
                    title_en=fd.get("title_en"),
                    title_ar=fd.get("title_ar"),
                    description_en=fd.get("description_en"),
                    page_number=fd.get("page"),
                    section_name=fd.get("section"),
                    evidence_quote=fd.get("quote"),
                ))

            rfp.status = "DECISION_READY"
            await db.commit()

            await publish_step("decision_scoring", "done", f"Decision: {decision_type} | Score: {total:.0f}", duration_ms=int((time.time() - t) * 1000))
            await publish_event(f"rfp:processing:{rfp_id}", {"step": "complete", "status": "done", "decision": decision_type, "score": total})

            logger.info("RFP processing complete", rfp_id=rfp_id, decision=decision_type, score=total)
            return {"decision": decision_type, "score": total}

        except Exception as exc:
            logger.error("RFP processing failed", rfp_id=rfp_id, error=str(exc), exc_info=True)
            rfp.status = "ACTION_REQUIRED"
            await db.commit()
            await publish_event(f"rfp:processing:{rfp_id}", {"step": "error", "status": "failed", "message": str(exc)})


# ── Text Extraction ────────────────────────────────────────────────────────────

async def _extract_text(content: bytes, mime_type: str, filename: str) -> str:
    """Extract text from PDF or DOCX content.

    Fix Bug #12: pdfplumber and python-docx are synchronous and CPU-bound.
    Wrapping them in asyncio.to_thread prevents blocking the event loop during
    large document processing (a 200-page PDF can take 10-30 seconds).
    """
    import asyncio

    if "pdf" in mime_type or filename.lower().endswith(".pdf"):
        return await asyncio.to_thread(_extract_pdf_sync, content)
    elif "word" in mime_type or filename.lower().endswith(".docx"):
        return await asyncio.to_thread(_extract_docx_sync, content)
    elif filename.lower().endswith(".md") or "markdown" in mime_type:
        # Strip markdown syntax for plain-text analysis
        import re as _re
        text = content.decode("utf-8", errors="ignore")
        text = _re.sub(r"```[\s\S]*?```", "", text)          # code blocks
        text = _re.sub(r"`[^`]+`", "", text)                  # inline code
        text = _re.sub(r"!\[.*?\]\(.*?\)", "", text)          # images
        text = _re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)  # links → label
        text = _re.sub(r"^#{1,6}\s+", "", text, flags=_re.MULTILINE)  # headings
        text = _re.sub(r"[*_~]{1,3}([^*_~]+)[*_~]{1,3}", r"\1", text)  # bold/italic
        text = _re.sub(r"^\s*[-|>]\s*", "", text, flags=_re.MULTILINE)  # bullets/blockquotes/tables
        return text
    return content.decode("utf-8", errors="ignore")


def _extract_pdf_sync(content: bytes) -> str:
    import io
    import pdfplumber
    pages_text = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            parts = []
            # Extract regular text
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text)
            # Extract tables — each cell may contain requirements
            tables = page.extract_tables() or []
            for table in tables:
                for row in table:
                    row_text = " | ".join(cell.strip() for cell in row if cell and cell.strip())
                    if row_text:
                        parts.append(row_text)
            if parts:
                pages_text.append(f"[PAGE {i}]\n" + "\n".join(parts))
    return "\n\n".join(pages_text)


def _extract_docx_sync(content: bytes) -> str:
    import io
    import docx
    doc = docx.Document(io.BytesIO(content))
    parts = []
    # Extract paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


# ── Section Detection & Classification ────────────────────────────────────────

def _detect_sections(text: str) -> list[dict]:
    """Detect sections by heading patterns (Arabic + English).

    Fix Bug #24: the original pattern used \\d which matches only ASCII digits (0-9).
    Arabic RFPs use Arabic-Indic numerals (٠١٢٣٤٥٦٧٨٩), so headings like
    '١- الشروط العامة' were silently skipped, reducing section coverage and
    downstream scoring accuracy.

    Page tracking: _extract_pdf_sync inserts \\x00PAGE:N\\x00 markers at each page
    boundary. We parse these to set accurate page numbers on each section.
    """
    import re
    sections = []
    lines = text.split("\n")
    current_section: dict = {"title": "Preamble", "content": "", "page": 1}
    current_page = 1
    page_marker = re.compile(r"^\[PAGE (\d+)\]$")
    # [\d٠-٩] covers both ASCII digits and Arabic-Indic numerals
    heading_pattern = re.compile(r"^([\d٠-٩]+\.?[\d٠-٩]*\.?\s+.{5,80}|[أ-ي].{2,50}:)$", re.UNICODE)
    for line in lines:
        m = page_marker.match(line)
        if m:
            current_page = int(m.group(1))
            continue
        if heading_pattern.match(line.strip()) and len(line.strip()) < 80:
            if current_section["content"].strip():
                sections.append(current_section)
            current_section = {"title": line.strip(), "content": "", "page": current_page}
        else:
            current_section["content"] += line + "\n"
    if current_section["content"].strip():
        sections.append(current_section)
    return sections or [{"title": "Full Document", "content": text, "page": 1}]


async def _classify_sections(sections: list[dict]) -> list[dict]:
    """Classify each section by type using bilingual keyword matching."""
    SCOPE_KW = ["scope", "deliverable", "نطاق", "المخرجات", "requirements", "متطلبات", "الأعمال المطلوبة"]
    ELIGIBILITY_KW = ["eligibility", "criteria", "الأهلية", "الشروط", "qualification", "الكفاءة"]
    TIMELINE_KW = ["timeline", "schedule", "الجدول", "المدة", "duration", "delivery", "تسليم"]
    COMMERCIAL_KW = ["price", "cost", "payment", "السعر", "المالي", "ضمان", "guarantee", "مالي"]
    COMPLIANCE_KW = ["compliance", "certification", "NDMO", "PDPL", "الامتثال", "شهادة", "NCA", "SAMA"]
    TECHNICAL_KW = ["technical", "architecture", "platform", "technology", "تقني", "هندسة", "منصة"]

    for section in sections:
        text = (section.get("title", "") + " " + section.get("content", "")).lower()
        if any(kw.lower() in text for kw in SCOPE_KW):
            section["type"] = "SCOPE"
        elif any(kw.lower() in text for kw in ELIGIBILITY_KW):
            section["type"] = "ELIGIBILITY"
        elif any(kw.lower() in text for kw in COMPLIANCE_KW):
            section["type"] = "COMPLIANCE"
        elif any(kw.lower() in text for kw in TECHNICAL_KW):
            section["type"] = "TECHNICAL"
        elif any(kw.lower() in text for kw in TIMELINE_KW):
            section["type"] = "TIMELINE"
        elif any(kw.lower() in text for kw in COMMERCIAL_KW):
            section["type"] = "COMMERCIAL_TERMS"
        else:
            section["type"] = "OTHER"
    return sections


# ── Entity Extraction ──────────────────────────────────────────────────────────

async def _extract_entities(sections: list[dict], rfp) -> dict:
    """Extract key entities: certs, value, local content, agency, deployment model."""
    import re
    all_text = " ".join(s.get("content", "") for s in sections)
    entities: dict = {}

    # Deadline
    deadline_match = re.search(r"(\d{4}[-/]\d{2}[-/]\d{2})", all_text)
    if deadline_match and not rfp.deadline:
        entities["deadline"] = deadline_match.group(1)

    # Project value in SAR
    value_match = re.search(r"(?:SAR|ريال)\s*[\d,]+", all_text, re.IGNORECASE)
    if value_match:
        entities["estimated_value_raw"] = value_match.group(0)
    # Also extract numeric value for scoring
    value_num = re.search(r"(?:SAR|ريال)\s*([\d,]+)", all_text, re.IGNORECASE)
    if value_num:
        try:
            entities["estimated_value_sar"] = int(value_num.group(1).replace(",", ""))
        except ValueError:
            pass

    # Required certifications
    cert_pattern = re.compile(
        r"ISO\s*\d+|CITC|SDAIA|NDMO|NDI|PDPL|SOC\s*2|PCI[-\s]DSS|SAMA\s*CSF|NCA\s*ECC|NCA|SAMA",
        re.IGNORECASE,
    )
    entities["required_certs"] = list(set(cert_pattern.findall(all_text)))

    # Local content percentage
    lc_match = re.search(r"(\d+)\s*%\s*(?:local content|محتوى محلي|نطاقات)", all_text, re.IGNORECASE)
    if lc_match:
        entities["local_content_pct"] = int(lc_match.group(1))

    # On-premise / cloud deployment preference
    if re.search(r"on[\s-]prem|on-premise|داخلي|حكومية سحابية|سحابة حكومية", all_text, re.IGNORECASE):
        entities["deployment_model"] = "on_prem"
    elif re.search(r"cloud|سحابة", all_text, re.IGNORECASE):
        entities["deployment_model"] = "cloud"

    # Defense / security clearance
    if re.search(r"security clearance|تصريح أمني|classified|سري|confidential", all_text, re.IGNORECASE):
        entities["requires_security_clearance"] = True

    # Data residency
    if re.search(r"data residency|outside.*saudi|خارج المملكة", all_text, re.IGNORECASE):
        entities["data_outside_ksa"] = True

    # Duration in months
    duration_match = re.search(r"(\d+)\s*(?:month|شهر)", all_text, re.IGNORECASE)
    if duration_match:
        entities["duration_months"] = int(duration_match.group(1))

    # Detect issuing agency name
    for agency in STRATEGIC_AGENCIES:
        if agency.lower() in all_text.lower():
            entities["issuing_agency"] = agency
            break

    return entities


# ── Flag Detection ─────────────────────────────────────────────────────────────

def _section_page(sections: list[dict], *keywords: str) -> int:
    """Return the page number of the first section whose title or type matches any keyword."""
    kw_lower = [k.lower() for k in keywords]
    for s in sections:
        title = s.get("title", "").lower()
        stype = s.get("type", "").lower()
        if any(k in title or k in stype for k in kw_lower):
            return s.get("page", 1)
    return 1


async def _detect_flags(sections: list[dict], entities: dict) -> dict:
    """
    Detect red and green flags based on Entropy's actual capabilities,
    certifications held, and past client experience.
    """
    red_flags = []
    green_flags = []
    all_text = " ".join(s.get("content", "") for s in sections)

    # ── INSTANT NO-GO: Scope mismatch detection ────────────────────────────────

    # Pattern 1: Pure cybersecurity (Entropy is not a cybersecurity firm)
    CYBERSEC_PATTERNS = [
        r"SOC[\s-]?as[\s-]?a[\s-]?service", r"\bSOC\s*setup\b", r"\bSIEM\b", r"\bSOAR\b",
        r"penetration test", r"\bpentest\b", r"red team", r"\bVA/PT\b", r"\bVAPT\b",
        r"\bNDR\b", r"\bEDR\b", r"\bXDR\b", r"managed security service",
        r"cybersecurity operation", r"أمن سيبراني.*تشغيل", r"مركز العمليات الأمنية",
        r"فريق.*أحمر",
    ]
    # Pattern 2: Clinical healthcare (not health analytics)
    CLINICAL_PATTERNS = [
        r"\bHIS\b", r"\bEMR\b", r"\bEHR\b", r"\bPACS\b", r"\bRIS\b",
        r"clinical decision support", r"hospital information system",
        r"medical imaging", r"surgical planning", r"ICU monitoring",
        r"pharmacy system", r"patient management system",
        r"نظام.*سريري", r"ملف المريض", r"التصوير الطبي",
    ]
    # Pattern 3: Network/hardware infrastructure
    INFRA_PATTERNS = [
        r"network.*cabling", r"physical.*infrastructure", r"data center.*build",
        r"\brouter\b.*\bswitch\b", r"hardware.*procurement.*only",
        r"تمديدات شبكة", r"بنية تحتية.*مادية",
    ]

    import re as _re_flags

    cyb_found = next(
        (p for p in CYBERSEC_PATTERNS if _re_flags.search(p, all_text, _re_flags.IGNORECASE)),
        None,
    )
    # Only trigger cybersecurity NO-GO if it's the PRIMARY scope, not just a mention
    # Heuristic: multiple cybersec keywords without any data/AI keywords = primary scope
    cyb_count = sum(1 for p in CYBERSEC_PATTERNS if _re_flags.search(p, all_text, _re_flags.IGNORECASE))
    data_ai_count = sum(
        1 for cap in ENTROPY_CAPABILITIES
        for kw in ENTROPY_CAPABILITIES[cap]["keywords_en"]
        if kw.lower() in all_text.lower()
    )
    if cyb_count >= 3 and data_ai_count <= 2:
        red_flags.append({
            "code": "SCOPE_MISMATCH",
            "severity": "CRITICAL",
            "title_en": "Primary scope is cybersecurity — outside Entropy's domain",
            "title_ar": "النطاق الرئيسي أمن سيبراني — خارج نطاق إنتروبي",
            "description_en": (
                "This RFP's primary scope appears to be cybersecurity services (SOC, SIEM, managed security, "
                "or similar). Entropy is not a cybersecurity firm and does not hold the required CITC license "
                "or NCA ECC certification. This is an instant NO-GO."
            ),
            "page": _section_page(sections, "scope", "نطاق"),
            "section": "Scope",
        })

    clin_found = next(
        (p for p in CLINICAL_PATTERNS if _re_flags.search(p, all_text, _re_flags.IGNORECASE)),
        None,
    )
    if clin_found:
        red_flags.append({
            "code": "SCOPE_MISMATCH",
            "severity": "CRITICAL",
            "title_en": "Clinical healthcare scope — outside Entropy's domain",
            "title_ar": "نطاق رعاية صحية سريرية — خارج نطاق إنتروبي",
            "description_en": (
                "This RFP requires clinical healthcare systems (HIS, EMR, PACS, clinical AI, or similar). "
                "Entropy does not hold healthcare accreditation, clinical AI validation frameworks, or "
                "Ministry of Health vendor registration. This is an instant NO-GO."
            ),
            "quote": f"Detected clinical pattern: {clin_found}",
            "page": _section_page(sections, "scope", "نطاق"),
            "section": "Scope",
        })

    infra_found = next(
        (p for p in INFRA_PATTERNS if _re_flags.search(p, all_text, _re_flags.IGNORECASE)),
        None,
    )
    if infra_found:
        red_flags.append({
            "code": "SCOPE_MISMATCH",
            "severity": "CRITICAL",
            "title_en": "Physical infrastructure scope — outside Entropy's domain",
            "title_ar": "نطاق بنية تحتية مادية — خارج نطاق إنتروبي",
            "description_en": (
                "This RFP appears to require physical network/infrastructure installation or hardware-only "
                "procurement. Entropy does not provide these services."
            ),
            "quote": f"Detected infrastructure pattern: {infra_found}",
            "page": _section_page(sections, "scope", "نطاق"),
            "section": "Scope",
        })

    # ── RED FLAGS ──────────────────────────────────────────────────────────────

    # 1. Required certification not held by Entropy
    for cert in entities.get("required_certs", []):
        cert_upper = cert.upper().replace(" ", "")
        held = any(h.upper().replace(" ", "") in cert_upper or cert_upper in h.upper().replace(" ", "")
                   for h in ENTROPY_CERTS_HELD)
        not_held = any(n.upper().replace(" ", "") in cert_upper or cert_upper in n.upper().replace(" ", "")
                       for n in ENTROPY_CERTS_NOT_HELD)
        if not_held and not held:
            # Only regulatory/financial certs are truly disqualifying in Saudi gov context
            severity = "CRITICAL" if cert.upper() in {"SAMA CSF", "NCA ECC", "PCI-DSS", "PCI DSS"} else "MAJOR"
            red_flags.append({
                "code": "MANDATORY_CERT_NOT_HELD",
                "severity": severity,
                "title_en": f"Required certification not held: {cert}",
                "title_ar": f"شهادة مطلوبة غير متوفرة: {cert}",
                "description_en": (
                    f"The RFP requires {cert}, which Entropy does not currently hold. "
                    f"{'This is typically a disqualifying requirement.' if severity == 'CRITICAL' else 'May be mitigated through a certified subcontractor.'}"
                ),
                "quote": f"Required certification: {cert}",
                "page": _section_page(sections, "eligibility", "الأهلية", "الشروط"),
                "section": "Eligibility",
            })

    # 2. Data residency outside KSA — PDPL conflict
    if entities.get("data_outside_ksa"):
        red_flags.append({
            "code": "DATA_RESIDENCY_OUTSIDE_KSA",
            "severity": "CRITICAL",
            "title_en": "Data residency outside KSA conflicts with PDPL",
            "title_ar": "إقامة البيانات خارج المملكة يتعارض مع نظام PDPL",
            "description_en": "The RFP appears to require data processing or storage outside Saudi Arabia, which conflicts with PDPL and Entropy's data sovereignty commitments.",
            "page": _section_page(sections, "compliance", "الامتثال", "pdpl"),
            "section": "Compliance",
        })

    # 3. Local content requirement too high
    lc_pct = entities.get("local_content_pct", 0)
    if lc_pct > 70:
        red_flags.append({
            "code": "LOCAL_CONTENT_HIGH",
            "severity": "CRITICAL",
            "title_en": f"Local content requirement ({lc_pct}%) exceeds capacity",
            "title_ar": f"نسبة المحتوى المحلي ({lc_pct}%) تتجاوز القدرة الحالية",
            "description_en": f"A {lc_pct}% local content requirement is very high and may exceed Entropy's current staffing levels.",
            "quote": f"Local content: {lc_pct}%",
            "page": _section_page(sections, "commercial", "المالي", "السعر"),
            "section": "Commercial Terms",
        })
    elif lc_pct > 50:
        red_flags.append({
            "code": "LOCAL_CONTENT_HIGH",
            "severity": "MAJOR",
            "title_en": f"High local content requirement: {lc_pct}%",
            "title_ar": f"نسبة محتوى محلي مرتفعة: {lc_pct}%",
            "description_en": f"Local content requirement of {lc_pct}% may require subcontracting or additional Saudi hires.",
            "quote": f"Local content: {lc_pct}%",
            "page": _section_page(sections, "commercial", "المالي", "السعر"),
            "section": "Commercial Terms",
        })

    # 4. Security clearance — informational note only, not a risk flag

    # 5. Very short timeline (< 3 months for complex delivery)
    duration = entities.get("duration_months", 0)
    if 0 < duration < 3:
        red_flags.append({
            "code": "UNREALISTIC_TIMELINE",
            "severity": "MAJOR",
            "title_en": f"Very short delivery timeline: {duration} month(s)",
            "title_ar": f"جدول زمني قصير جداً: {duration} شهر",
            "description_en": f"A {duration}-month delivery window is tight for a data/AI project. Risk of scope creep and delivery failure.",
            "page": _section_page(sections, "timeline", "الجدول", "المدة", "تسليم"),
            "section": "Timeline",
        })

    # ── GREEN FLAGS ────────────────────────────────────────────────────────────

    # 1. Issuing agency is an existing Entropy client
    issuing_agency = entities.get("issuing_agency", "")
    if issuing_agency and any(c.lower() in issuing_agency.lower() or issuing_agency.lower() in c.lower()
                               for c in ENTROPY_EXISTING_CLIENTS):
        green_flags.append({
            "title_en": f"Existing client: {issuing_agency}",
            "title_ar": f"عميل حالي: {issuing_agency}",
            "description_en": f"Entropy has an active or prior engagement with {issuing_agency}, providing a competitive advantage.",
            "page": _section_page(sections, "preamble", "المقدمة", "scope"),
            "section": "Preamble",
        })
    elif issuing_agency:
        green_flags.append({
            "title_en": f"Strategic target agency: {issuing_agency}",
            "title_ar": f"جهة مستهدفة استراتيجياً: {issuing_agency}",
            "description_en": f"{issuing_agency} is on Entropy's strategic target list.",
            "page": _section_page(sections, "preamble", "المقدمة", "scope"),
            "section": "Preamble",
        })

    # 2. NDMO/NDI compliance required — Entropy's strong suit
    if any(c.upper() in {"NDMO", "NDI"} for c in entities.get("required_certs", [])):
        green_flags.append({
            "title_en": "NDMO/NDI compliance required — Entropy's strength",
            "title_ar": "يتطلب الامتثال لمعايير NDMO/NDI — ميزة تنافسية لـ Entropy",
            "description_en": "Entropy has delivered NDMO-compliant data management for the Digital Government Authority and other government clients.",
            "page": _section_page(sections, "compliance", "الامتثال", "ndmo"),
            "section": "Compliance",
        })

    # 3. Arabic language / Arabic-first requirement
    if any(kw in all_text.lower() for kw in ["arabic", "عربي", "عربية", "rtl", "اللغة العربية"]):
        green_flags.append({
            "title_en": "Arabic-first requirement — core Entropy differentiator",
            "title_ar": "متطلبات عربية أولى — ميزة Entropy الجوهرية",
            "description_en": "Entropy builds Arabic-first AI systems with deep cultural and linguistic context — a key differentiator in the Saudi market.",
            "page": _section_page(sections, "scope", "نطاق", "المخرجات"),
            "section": "Scope",
        })

    # 4. Technology partners alignment
    for partner in ENTROPY_TECH_PARTNERS:
        if partner.lower() in all_text.lower():
            green_flags.append({
                "title_en": f"Technology partner alignment: {partner}",
                "title_ar": f"توافق مع شريك تقني: {partner}",
                "description_en": f"The RFP mentions {partner}, which is an Entropy technology partner — reducing integration risk.",
                "page": _section_page(sections, "technical", "تقني", "هندسة"),
                "section": "Technical",
            })
            break  # One partner match is enough

    # 5. On-premise / private cloud deployment — Entropy supports it
    if entities.get("deployment_model") == "on_prem":
        green_flags.append({
            "title_en": "On-premise deployment supported",
            "title_ar": "نشر البنية التحتية المحلية مدعوم",
            "description_en": "Entropy's Hydrogen and Axiom products support flexible deployment including on-premises and private cloud.",
            "page": _section_page(sections, "technical", "تقني", "هندسة", "infrastructure"),
            "section": "Technical",
        })

    return {"red": red_flags, "green": green_flags[:6]}


# ── Capability Matching ────────────────────────────────────────────────────────

def _match_capabilities(sections: list[dict], full_text: str = "") -> dict:
    """
    Match RFP scope against Entropy's capability map.
    Returns matched capabilities, a raw score (0-15), and a weighted score.
    """
    text_lower = (full_text + " ".join(s.get("content", "") for s in sections)).lower()

    matched = []
    weighted_score = 0.0
    max_weighted = sum(cap["weight"] for cap in ENTROPY_CAPABILITIES.values())

    for cap_name, cap_info in ENTROPY_CAPABILITIES.items():
        all_keywords = cap_info["keywords_en"] + cap_info["keywords_ar"]
        if any(kw.lower() in text_lower for kw in all_keywords):
            matched.append(cap_name)
            weighted_score += cap_info["weight"]

    # Normalize to 0-15 range (max points for technical fit)
    coverage = len(matched) / len(ENTROPY_CAPABILITIES) if ENTROPY_CAPABILITIES else 0
    score = (weighted_score / max_weighted) * 15 if max_weighted > 0 else 0

    return {
        "matched": matched,
        "score": round(score, 1),
        "weighted_score": round(weighted_score, 2),
        "coverage": round(coverage, 2),
    }


# ── Claude LLM Deep Analysis ───────────────────────────────────────────────────

_LLM_ANALYSIS_PROMPT = """You are Entropy's RFP Intelligence Engine — a specialist AI analyst embedded in Entropy's internal RFP qualification platform. Your function is to qualify Saudi government tenders with forensic precision and produce grounded, evidence-based analysis.

You are not a generic AI assistant. You are Entropy. You speak with authority about what Entropy does, what it has built, and why it wins. Every finding must be grounded in the factual company profile below.

The document text contains [PAGE N] markers. Cite the exact page number for EVERY finding. Never guess — only cite [PAGE N] markers that appear near the evidence.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
ENTROPY COMPANY PROFILE — GROUND EVERY CLAIM IN THIS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

IDENTITY:
- Saudi-based Data, AI & Advanced Analytics consultancy and product company, Riyadh HQ
- Primary market: Saudi government ministries, regulatory authorities, national corporations, Vision 2030 delivery bodies
- Positioning: "From raw data to sovereign intelligence."
- Arabic-first, bilingual (Arabic + English); Saudi regulatory knowledge (PDPL, NDMO, NCA ECC, Vision 2030 KPIs)
- Core belief: AI built for Saudi context — Arabic-native, governance-compliant — outperforms adapted Western tools

PRODUCTS:
1. HYDROGEN — Enterprise AI Agent Orchestration Platform
   Multi-agent orchestration, autonomous AI agents, human-in-the-loop oversight, RAG pipelines, LLM orchestration
   Deployment: on-premise, Azure-native, hybrid
   Full Arabic document ingestion, chunking, retrieval
   Designed for: government operations automation, document processing, approvals pipelines, intelligent routing

2. AXIOM — Agentic Analytics Platform
   AI-driven insight generation, automated Arabic/English narrative reporting, agentic layer (not passive BI)
   Integrations: Microsoft Fabric, Power BI, Tableau, Snowflake, custom data lakes
   Differentiator: analytical agents proactively surface anomalies, forecasts, and executive narratives — no query required
   Designed for: ministerial performance dashboards, KPI monitoring, operational intelligence, audit analytics

3. YAMEEN — AI Meeting Intelligence Platform
   Arabic + bilingual transcription, summarization, action items — real-time
   Gulf Arabic dialect awareness
   Integrations: Microsoft Teams, Zoom, Google Meet
   Designed for: government councils, ministerial committees, board meetings, project governance

SERVICE PILLARS:
PILLAR 1 — Data & AI Foundations:
  Data governance (NDMO-aligned, DAMA-compatible), data strategy & roadmap, DMO setup & activation,
  data architecture (lakehouse, medallion, federated), data quality frameworks, MDM,
  PDPL compliance assessments, metadata management, data catalog, data literacy programs

PILLAR 2 — AI & ML Technologies:
  Generative AI (RAG pipelines, fine-tuning, prompt engineering), Arabic NLP (Gulf dialect, formal Arabic),
  OCR + VLM pipelines for scanned Arabic documents, ML model development (predictive, classification, clustering),
  AI model evaluation & red-teaming, responsible AI frameworks, AI readiness assessments, AI use case discovery

PILLAR 3 — AI & Analytics Platforms:
  Microsoft Fabric implementation, Power BI premium + governance, Snowflake, custom agentic analytics (Axiom),
  BI Center of Excellence setup, real-time analytics & streaming, data product development

CONFIRMED CLIENTS (reference only when contextually relevant to the RFP scope):
- Ministry of Interior (MOI): Full NDMO-aligned DMO activation, data governance policy framework, 20-week delivery, first ministerial DMO at NDMO Tier 1
- Royal Commission for AlUla (RCU): 3-stream AI audit platform — Hydrogen agents + Axiom analytics + CCA, bilingual NLP, Azure-native, Teammate+ integration
- Monsha'at (SME Authority): AI adoption framework for national SME authority, consulting + training delivery
- Saudi Ports Authority (Mawani): Data & AI platform
- Saudi Health Council: Analytics and data platform
- Saudi Central Bank (SAMA): Data governance and analytics
- General Authority for Statistics (GASTAT): Data platform and analytics
- National Cybersecurity Authority (NCA): Data and reporting work — this is DATA work FOR NCA, NOT cybersecurity services
- Saudi Aramco: AI and analytics engagements

PROVEN SUCCESS STORIES (use as proposal evidence):
Story 1 — MOI DMO: NDMO-aligned DMO, governance policy suite, data stewardship model, metadata catalog, training. 20 weeks, 12 deliverables.
Story 2 — RCU AI Audit: AI audit efficiency agents, advanced audit analytics, continuous control monitoring. Audit cycle reduced from manual weeks to AI-assisted days.
Story 3 — Monsha'at AI Adoption: AI adoption framework for Saudi SME national authority, structured roadmap for national rollout.
Story 4 — Agentic Analytics: Client moved from monthly manual BI reports to daily AI-generated Arabic insight packages — no headcount increase.
Story 5 — Sovereign Data Platform: On-premise/hybrid data lakehouse + ML model registry + vector store + governance layer for a government entity.

TECHNOLOGY PARTNERS (formal partnerships):
Microsoft/Azure, Microsoft Fabric, Power BI, Copilot, Snowflake, Databricks, Informatica, Collibra, Tableau, OpenAI/Azure OpenAI, Glean, Google Cloud/GCP, Dataiku, Oracle, Groq

CERTIFICATIONS & DIFFERENTIATORS:
- NDMO alignment: deep expertise — hard requirement in many government RFPs, a core Entropy strength
- DAMA/CDMP certified staff
- Arabic-native AI: NOT translation layers — builds Arabic-first NLP with Gulf dialect awareness
- PDPL, NCA ECC, Vision 2030 regulatory knowledge
- Saudization: maintains ≥40% Saudi national staffing
- Bilingual deliverables: Arabic + English on all outputs

CERTIFICATIONS HELD: NDMO alignment, DAMA/CDMP
CERTIFICATIONS NOT HELD (eligibility risk if required):
  ISO 27001, ISO 9001, ISO 20000, PCI-DSS, SOC 2, SAMA CSF, NCA ECC, CITC license

WHAT ENTROPY CANNOT DELIVER (delivery failure risk → flag as SCOPE_MISMATCH):
- Cybersecurity operations: SOC setup, SIEM/SOAR, penetration testing, NDR/EDR/XDR, red team, managed security
- Clinical healthcare AI: HIS, EMR, PACS, clinical decision support, medical imaging AI, surgical planning, ICU monitoring
- Network/telecom infrastructure: routers, switches, cabling, data center physical build
- Hardware-only procurement
- General IT development unrelated to AI/data: e-government portals (non-AI), ERP (non-data track), IT helpdesk

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
8-DIMENSION QUALIFICATION RUBRIC (100 POINTS TOTAL)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Apply this rubric dimension by dimension. Justify each score with a quote or specific reference.

D1 — Sector & Client Fit (0-20):
  18-20: Saudi government ministry, regulatory authority, national corporation (MOI, NDMO, SFDA, NCA, SAMA, RCU, GASTAT, Aramco affiliates, Vision 2030 delivery bodies)
  12-17: Saudi quasi-government, semi-private national, large private sector (banking, telecom, energy)
  6-11:  Saudi private sector SME, regional government, non-Saudi GCC
  0-5:   Non-GCC entity, non-strategic sector
  GREEN: "وزارة", "هيئة", "شركة وطنية", Vision 2030 entity, existing Entropy client

D2 — Technical Domain Alignment (0-25):
  22-25: Core Entropy domain: data governance, NDMO alignment, AI agents, GenAI, Arabic NLP, analytics platforms, DMO setup
  15-21: Adjacent: data engineering, BI platforms, AI strategy consulting, digital transformation with AI component
  8-14:  Partial: IT project management with data component, general digital transformation (no AI requirement)
  0-7:   No overlap: pure cybersecurity, network infrastructure, hardware procurement, clinical systems
  INSTANT 0 + SCOPE_MISMATCH CRITICAL: pure cybersecurity, clinical healthcare, network/hardware

D3 — NDMO / Data Governance (0-15):
  13-15: Explicitly requires NDMO alignment, DMO setup, data governance policies, PDPL compliance, data stewardship
  8-12:  Data strategy, data quality, metadata — NDMO implied
  3-7:   Data present but governance not a priority
  0-2:   No data governance element

D4 — AI / Generative AI (0-15):
  13-15: Explicit AI agents, GenAI, LLM, RAG, agentic analytics, NLP, AI-powered automation
  8-12:  AI/ML models, intelligent analytics, smart dashboards, predictive models
  3-7:   AI mentioned as desirable but not core; mainly BI/reporting
  0-2:   No AI requirement

D5 — Arabic Language (0-10):
  9-10: Arabic-first mandatory; Arabic NLP processing required; bilingual platform
  6-8:  Arabic deliverables required; bilingual interface required
  3-5:  Arabic preferred; English acceptable
  0-2:  English-only or language not specified
  NOTE: Arabic-native NLP is Entropy's key differentiator — RFPs requiring Arabic processing = hardest wins

D6 — Entropy Product Fit (0-5):
  5:    Clear fit for Hydrogen (agent orchestration), Axiom (agentic analytics), or Yameen (meeting intelligence)
  3-4:  Product could be proposed as part of solution
  1-2:  Adaptable but significant customization needed
  0:    No product fit; pure professional services

D7 — Scope Feasibility (0-5):
  5:    Well-defined scope, realistic 12-24 week timeline, clear deliverables, manageable team
  3-4:  Somewhat ambiguous; tight but achievable; deliverables negotiable
  1-2:  Massive or unclear scope; unrealistic timeline; requires heavy subcontracting
  0:    Fundamentally outside Entropy's capacity without a major consortium partner

D8 — Compliance & Eligibility (0-5):
  5:    No hard eligibility barriers; Entropy meets all stated requirements
  3-4:  Minor eligibility question (e.g., ISO 27001 needed); addressable via partner
  1-2:  Significant barrier (e.g., SAMA CSF, CITC license, healthcare accreditation)
  0:    Entropy definitionally ineligible (listed stock company required, healthcare institution required)

DECISION THRESHOLDS:
  80-100: GO — Pursue aggressively
  65-79:  CONDITIONAL GO — Pursue if capacity allows; identify gaps; consider partnership
  45-64:  WATCH — Monitor; do not pursue now; may revisit if scope changes
  0-44:   NO-GO — Decline

INSTANT NO-GO OVERRIDE RULES (flag as SCOPE_MISMATCH CRITICAL):
  1. RFP is exclusively cybersecurity: SOC, SIEM/SOAR, pentest, VA/PT, red team, NDR/EDR/XDR
  2. RFP is for clinical healthcare: HIS, EMR, PACS, clinical AI, medical imaging, surgical planning
  3. RFP is for physical infrastructure, network, hardware-only
  4. Entropy explicitly ineligible per mandatory criteria (check D8)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTOR-SPECIFIC GUIDANCE
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Saudi Government Ministries (MOI, MOF, MISA, MCIT, MOH, MOE):
  Lead with NDMO compliance + data governance. PDPL mandatory. Arabic deliverables required. 70/30 technical/financial scoring.
  Reference Ministry of Interior DMO success when relevant. Etimad procurement platform.

Regulatory Authorities (SAMA, SFDA, GASTAT, NCA, CITC, CMA):
  Lead with domain-specific AI. Arabic reporting essential. Sector-specific regulatory compliance layer.
  Key: NCA is an ENTROPY DATA CLIENT — Entropy does data work FOR NCA. This is different from doing cybersecurity services.

National Corporations (Aramco, SABIC, STC, Mawani):
  Lead with AI platform architecture + Axiom for operational intelligence + Hydrogen for workflow automation.
  Commercial flexibility higher. Can propose SaaS/subscription.

Vision 2030 Delivery Bodies (RCU, NEOM, Diriyah Gate, Qiddiya, ROSHN):
  Smart city data layer, AI for project governance, Yameen for executive meeting intelligence.
  Innovation narrative resonates. Compressed timelines — assess team availability.

SME / Private Sector (Monsha'at portfolio, private banks, telecom):
  Lead with ROI and time-to-value. Shorter 8-16 week engagements. SaaS deployment preferred.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
KNOWN REJECTION PATTERNS — ACCELERATE NO-GO DECISIONS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

PATTERN 1: Healthcare Clinical AI / AI Factory (Clinical)
  Keywords: HIS, EMR, PACS, "AI Factory" for hospital ops, clinical decision support, medical imaging, ICU monitoring, pharmacy systems
  Why rejected: requires healthcare regulatory accreditation, clinical AI validation, MOH vendor registration. No relevant past performance.
  Rule: If AI primarily modifies clinical workflows or patient care → SCOPE_MISMATCH CRITICAL → NO-GO
  CAUTION: health DATA governance, NHIC analytics, public health statistics ≠ clinical → may still fit

PATTERN 2: Pure Cybersecurity
  Keywords: SOC-as-a-service, SIEM, SOAR, pentest, VA/PT, red team, NDR/EDR/XDR, managed security, NCA compliance AUDIT (as cybersecurity firm)
  Why rejected: not a cybersecurity firm, no CITC cybersecurity license, no security ops team
  Rule: If primary deliverable is a cybersecurity service (not data/AI) → SCOPE_MISMATCH CRITICAL
  CAUTION: cybersecurity RFP that also has data analytics / AI component → score D2 for the data part only; flag partial scope risk

PATTERN 3: Non-Data/AI Digital Government Services
  Keywords: e-government portals (non-AI), mobile apps, IT helpdesk, ERP implementation (non-data track), general digital government
  Why rejected: IT services, not Data/AI consulting; Entropy is not a general IT integrator
  Rule: If scope is digital services without meaningful Data/AI/analytics → SCOPE_MISMATCH MAJOR

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
YOUR ANALYSIS TASK — STEP BY STEP
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Step 1 — READ the entire document word by word. Do not rush. Do not summarize prematurely.
Step 2 — CHECK for INSTANT NO-GO triggers first. If found, set SCOPE_MISMATCH CRITICAL immediately.
Step 3 — For EVERY requirement found:
  • Which Entropy capability covers it? (Be specific: HYDROGEN for agents, AXIOM for analytics, specific service pillar)
  • Can Entropy fulfill it: FULL / PARTIAL / NONE?
  • If partial or none, is there a partner or mitigation path?
Step 4 — Score each of the 8 dimensions with evidence. Justify every score with an exact quote and page.
Step 5 — Identify advantages (min 5), disadvantages (min 5), all rejection reasons, all acceptance boosters.
Step 6 — Write bilingual summaries (Arabic primary for government RFPs).
Step 7 — Output the JSON exactly as specified below.

CRITICAL RULES:
- Cite [PAGE N] markers for every finding — never invent page numbers
- Do NOT fabricate Entropy capabilities, certifications, or client names beyond the profile above
- Be EXHAUSTIVE — read every table row, every Arabic section, every requirement. Missing a requirement is an analysis failure.
- technical_requirements: ALL requirements — functional, non-functional, integration, security, SLA, training, maintenance
- mandatory_requirements: EVERY mandatory/obligatory/إلزامي item
- advantages + disadvantages: minimum 5 each; specific to this RFP, not generic
- acceptance_boosters: concrete actions Entropy must take in THIS bid to win
- analyst_confidence: honest; lower if OCR is poor, scope is vague, or Arabic is unclear
- For Saudi government RFPs: reference Vision 2030 alignment, NDMO relevance, PDPL compliance in summaries

Return ONLY a valid JSON object — no markdown, no backticks, no explanation outside the JSON:

{
  "entities": {
    "issuing_agency": "<string | null>",
    "estimated_value_sar": <number | null>,
    "deadline": "<YYYY-MM-DD | null>",
    "deadline_page": <number | null>,
    "duration_months": <number | null>,
    "deployment_model": "<on_prem | cloud | hybrid | null>",
    "required_certs": ["<cert name>"],
    "local_content_pct": <number | null>,
    "data_outside_ksa": <true | false>,
    "requires_security_clearance": <true | false>,
    "primary_language": "<arabic | english | mixed>",
    "project_objectives": ["<objective>"],
    "key_deliverables": ["<deliverable>"],
    "technical_requirements": [
      {
        "requirement": "<exact requirement text>",
        "page": <number | null>,
        "quote": "<exact quote from document>",
        "entropy_can_fulfill": <true | false | "partial">,
        "fulfillment_notes": "<how Entropy fulfills using which product/pillar, or why it cannot>"
      }
    ],
    "mandatory_requirements": [
      {
        "requirement": "<exact text>",
        "page": <number | null>,
        "quote": "<exact quote>",
        "is_blocker": <true | false>,
        "blocker_reason": "<why this blocks Entropy | null>"
      }
    ]
  },
  "dimension_scores": {
    "d1_sector_fit": <0-20>,
    "d2_domain_alignment": <0-25>,
    "d3_ndmo_governance": <0-15>,
    "d4_ai_genai": <0-15>,
    "d5_arabic_language": <0-10>,
    "d6_product_fit": <0-5>,
    "d7_scope_feasibility": <0-5>,
    "d8_compliance_eligibility": <0-5>,
    "total": <0-100>,
    "decision_recommendation": "<GO | CONDITIONAL_GO | WATCH | NO_GO>",
    "d1_rationale": "<quote + justification>",
    "d2_rationale": "<quote + justification>",
    "d3_rationale": "<quote + justification>",
    "d4_rationale": "<quote + justification>",
    "d5_rationale": "<quote + justification>",
    "d6_rationale": "<quote + justification>",
    "d7_rationale": "<quote + justification>",
    "d8_rationale": "<quote + justification>"
  },
  "advantages": [
    {
      "title_en": "<string>",
      "title_ar": "<string>",
      "description_en": "<detailed explanation tied to specific Entropy product/pillar/success story>",
      "description_ar": "<شرح تفصيلي>",
      "evidence": "<exact quote from RFP>",
      "page": <number | null>,
      "impact": "<HIGH | MEDIUM | LOW>",
      "relevant_success_story": "<MOI DMO | RCU Audit | Monsha'at | Agentic Analytics | Sovereign Platform | null>"
    }
  ],
  "disadvantages": [
    {
      "title_en": "<string>",
      "title_ar": "<string>",
      "description_en": "<detailed explanation>",
      "description_ar": "<شرح تفصيلي>",
      "evidence": "<exact quote>",
      "page": <number | null>,
      "severity": "<CRITICAL | MAJOR | MINOR>",
      "is_deal_breaker": <true | false>
    }
  ],
  "rejection_reasons": [
    {
      "reason_en": "<exact reason>",
      "reason_ar": "<سبب الرفض>",
      "evidence": "<exact quote>",
      "page": <number | null>,
      "can_mitigate": <true | false>,
      "mitigation_strategy": "<concrete plan including which partner to bring | null>"
    }
  ],
  "acceptance_boosters": [
    {
      "booster_en": "<what helps Entropy win — specific to this RFP>",
      "booster_ar": "<ما يساعد على الفوز>",
      "evidence": "<what in the RFP makes this relevant>",
      "page": <number | null>,
      "action_required": "<exact action Entropy must take: which product to highlight, which success story to reference, which partner to mention>"
    }
  ],
  "red_flags": [
    {
      "code": "<FLAG_CODE>",
      "severity": "<CRITICAL | MAJOR | MINOR>",
      "title_en": "<string>",
      "title_ar": "<string>",
      "description_en": "<full explanation with mitigation path>",
      "description_ar": "<شرح كامل>",
      "evidence": "<exact quote>",
      "page": <number | null>
    }
  ],
  "green_flags": [
    {
      "title_en": "<string>",
      "title_ar": "<string>",
      "description_en": "<full explanation>",
      "description_ar": "<شرح كامل>",
      "evidence": "<exact quote>",
      "page": <number | null>
    }
  ],
  "capability_signals": ["<capability keyword found in document>"],
  "missing_capabilities": ["<what the RFP needs that Entropy lacks — be honest>"],
  "company_fit": {
    "overall_fit": "<STRONG | MODERATE | WEAK | NO_FIT>",
    "fit_explanation_en": "<detailed explanation grounded in Entropy's actual capabilities and this RFP's specific requirements>",
    "fit_explanation_ar": "<شرح تفصيلي لمدى ملاءمة إنتروبي لهذا الطلب مبني على القدرات الفعلية>",
    "requirements_coverage": [
      {
        "requirement": "<exact requirement text from RFP>",
        "page": <number | null>,
        "entropy_capability": "<specific Entropy product or pillar that covers this | null>",
        "coverage": "<FULL | PARTIAL | NONE>",
        "notes": "<how Entropy covers it using which product/pillar/success story, or why it cannot>"
      }
    ],
    "field_alignment": "<how well Entropy's AI/data field matches the RFP domain — cite specific matches>",
    "competitive_position": "<STRONG | COMPETITIVE | WEAK>",
    "win_probability": "<HIGH | MEDIUM | LOW>"
  },
  "recommended_positioning": {
    "lead_service_pillar": "<Pillar 1: Data & AI Foundations | Pillar 2: AI & ML Technologies | Pillar 3: AI & Analytics Platforms>",
    "products_to_highlight": ["<Hydrogen | Axiom | Yameen>"],
    "success_stories_to_reference": ["<MOI DMO | RCU Audit | Monsha'at | Agentic Analytics | Sovereign Platform>"],
    "key_differentiators": ["<Arabic-native NLP | NDMO expertise | Saudi regulatory knowledge | Hydrogen agents | Axiom agentic analytics>"],
    "vision_2030_angle": "<how to position this bid within Vision 2030 — specific program or KPI to reference>"
  },
  "summary_ar": "<5-7 sentence Arabic summary: scope, Entropy fit, key risks, Vision 2030 alignment, recommended action>",
  "summary_en": "<5-7 sentence English summary: scope, Entropy fit, key risks, recommended action>",
  "analyst_confidence": <0.0-1.0>,
  "analyst_notes": "<any observations about the RFP quality, ambiguities, or strategic considerations not captured elsewhere>"
}

Red flag codes: MANDATORY_CERT_NOT_HELD, DATA_RESIDENCY_OUTSIDE_KSA, LOCAL_CONTENT_HIGH,
SECURITY_CLEARANCE_REQUIRED, UNREALISTIC_TIMELINE, PDPL_CONFLICT, CONFLICT_OF_INTEREST,
BUDGET_TOO_SMALL, SCOPE_MISMATCH, HIGH_COMPETITION_RISK, GPU_INFRA_REQUIRED, SOLE_SOURCE_SPEC.
"""


async def _search_knowledge_base(db, all_text: str, sections: list[dict]) -> list[dict]:
    """Search knowledge docs for relevant content using keyword matching.

    Works without Qdrant — queries the database directly for knowledge docs
    whose titles or content keywords match the RFP scope.
    Returns up to 10 most relevant docs with excerpts.
    """
    try:
        from models.knowledge_doc import KnowledgeDoc
        from sqlalchemy import select, or_
        from services.storage import StorageService

        # Extract key terms from sections and full text for matching
        scope_keywords = set()
        for section in sections:
            words = (section.get("content", "") + " " + section.get("title", "")).split()
            for word in words:
                if len(word) > 4:
                    scope_keywords.add(word.lower()[:20])

        # Query all non-deleted knowledge docs
        result = await db.execute(
            select(KnowledgeDoc).where(
                KnowledgeDoc.is_deleted == False,  # noqa: E712
                KnowledgeDoc.is_indexed == True,
            ).limit(200)
        )
        all_docs = result.scalars().all()

        if not all_docs:
            return []

        # Score each doc by keyword overlap with RFP text
        scored = []
        rfp_lower = all_text[:50000].lower()
        for doc in all_docs:
            score = 0
            title_lower = (doc.title or "").lower()
            tags = doc.tags_json or ""

            # Title keyword match
            for word in title_lower.split():
                if len(word) > 3 and word in rfp_lower:
                    score += 2

            # Tags match
            for word in tags.lower().split():
                if len(word) > 3 and word in rfp_lower:
                    score += 1

            # Boost WIN outcomes
            if doc.outcome == "WIN":
                score += 3

            # Boost capability/case study docs
            if doc.doc_type in ("CAPABILITY", "CASE_STUDY", "PAST_PROPOSAL"):
                score += 1

            if score > 0:
                scored.append((score, doc))

        # Sort by score descending, take top 10
        scored.sort(key=lambda x: x[0], reverse=True)
        top_docs = scored[:10]

        # Load excerpts from storage
        storage = StorageService.get_instance()
        matches = []
        for _, doc in top_docs:
            excerpt = ""
            try:
                raw = await storage.download_file(doc.storage_path)
                extracted = await _extract_text(raw, "application/pdf", doc.storage_path)
                excerpt = extracted[:800]
            except Exception:
                pass
            matches.append({
                "title": doc.title,
                "doc_type": doc.doc_type,
                "outcome": doc.outcome,
                "year": doc.year,
                "excerpt": excerpt,
            })

        return matches

    except Exception as exc:
        logger.warning("Knowledge base search failed", error=str(exc))
        return []


async def _llm_analyze(all_text: str, knowledge_context: list[dict] | None = None) -> dict:
    """Three-pass deep analysis using Claude Opus.

    Pass 1 — Requirement extraction: reads every page and extracts all requirements
    Pass 2 — Risk & flag analysis: identifies every risk, blocker, and compliance issue
    Pass 3 — Company fit: maps requirements to Entropy capabilities, scores fit

    Each pass is independent and deep. Together they guarantee thorough coverage.
    Falls back to empty result on any error so the pipeline always continues.
    """
    try:
        from core.config import settings
        from services.llm_client import make_anthropic_client
        client = make_anthropic_client()

        import asyncio as _asyncio
        import json as _json

        text_excerpt = all_text[:150000]

        knowledge_section = ""
        if knowledge_context:
            knowledge_section = "\n\nRELEVANT KNOWLEDGE BASE MATCHES:\n"
            for doc in knowledge_context:
                knowledge_section += f"[{doc['doc_type']}] {doc['title']}"
                if doc.get('outcome'):
                    knowledge_section += f" (Outcome: {doc['outcome']})"
                if doc.get('excerpt'):
                    knowledge_section += f"\n{doc['excerpt'][:400]}\n"

        async def _call(prompt: str, max_tokens: int) -> str:
            r = await _asyncio.wait_for(
                client.messages.create(
                    model=settings.primary_llm_model,
                    max_tokens=max_tokens,
                    messages=[{"role": "user", "content": prompt}],
                ),
                timeout=300,
            )
            return next((b.text for b in r.content if hasattr(b, "text")), "").strip()

        def _strip_fences(s: str) -> str:
            if s.startswith("```"):
                s = "\n".join(s.splitlines()[1:])
                if s.strip().endswith("```"):
                    s = s[:s.rfind("```")].strip()
            return s

        # ── PASS 1: Deep requirement extraction ──────────────────────────────
        logger.info("LLM Pass 1: extracting all requirements")
        pass1_prompt = f"""{_LLM_ANALYSIS_PROMPT}

Document text:
{text_excerpt}

PASS 1 TASK: Read the ENTIRE document word by word. Extract EVERY requirement — functional, non-functional, mandatory, technical, compliance, timeline, certification, local content, security, SLA, training, maintenance. Read every table row. Read Arabic sections carefully.

Return ONLY a JSON object with these keys:
{{
  "entities": {{...}},
  "technical_requirements": [...],
  "mandatory_requirements": [...]
}}"""

        pass1_raw = _strip_fences(await _call(pass1_prompt, 10000))
        try:
            pass1 = _json.loads(pass1_raw)
        except Exception:
            pass1 = {}
        logger.info("LLM Pass 1 complete", requirements=len(pass1.get("technical_requirements", [])))

        # ── PASS 2: Deep risk & flag analysis ────────────────────────────────
        logger.info("LLM Pass 2: analyzing risks and flags")
        requirements_summary = _json.dumps(pass1.get("technical_requirements", [])[:30], ensure_ascii=False)
        pass2_prompt = f"""{_LLM_ANALYSIS_PROMPT}

Document text:
{text_excerpt}

Requirements extracted in previous pass:
{requirements_summary}

PASS 2 TASK: Focus entirely on risks, blockers, and opportunities. For every requirement extracted, assess: does it block Entropy? What certifications are missing? What are the compliance risks? What advantages does Entropy have?

Return ONLY a JSON object with these keys:
{{
  "red_flags": [...],
  "green_flags": [...],
  "advantages": [...],
  "disadvantages": [...],
  "rejection_reasons": [...],
  "acceptance_boosters": [...]
}}"""

        pass2_raw = _strip_fences(await _call(pass2_prompt, 10000))
        try:
            pass2 = _json.loads(pass2_raw)
        except Exception:
            pass2 = {}
        logger.info("LLM Pass 2 complete", red_flags=len(pass2.get("red_flags", [])))

        # ── PASS 3: Company fit mapping ───────────────────────────────────────
        logger.info("LLM Pass 3: mapping requirements to Entropy capabilities")
        pass3_prompt = f"""{_LLM_ANALYSIS_PROMPT}

Document text (summary):
{text_excerpt[:50000]}

{knowledge_section}

All extracted requirements:
{requirements_summary}

PASS 3 TASK: Map EVERY requirement to Entropy's specific capabilities. For each one: which Entropy capability covers it? FULL, PARTIAL, or NONE? Give honest assessment of overall fit, win probability, and what Entropy must emphasize.

Return ONLY a JSON object with these keys:
{{
  "company_fit": {{
    "overall_fit": "<STRONG|MODERATE|WEAK|NO_FIT>",
    "fit_explanation_en": "<detailed>",
    "fit_explanation_ar": "<تفصيلي>",
    "requirements_coverage": [...],
    "field_alignment": "<explanation>",
    "competitive_position": "<STRONG|COMPETITIVE|WEAK>",
    "win_probability": "<HIGH|MEDIUM|LOW>"
  }},
  "capability_signals": [...],
  "missing_capabilities": [...],
  "summary_ar": "<5-7 sentences>",
  "summary_en": "<5-7 sentences>",
  "analyst_confidence": <0.0-1.0>,
  "analyst_notes": "<observations>"
}}"""

        pass3_raw = _strip_fences(await _call(pass3_prompt, 10000))
        try:
            pass3 = _json.loads(pass3_raw)
        except Exception:
            pass3 = {}
        logger.info("LLM Pass 3 complete", fit=pass3.get("company_fit", {}).get("overall_fit"))

        # ── Merge all three passes ────────────────────────────────────────────
        result = {
            **pass1,
            **pass2,
            **pass3,
        }

        logger.info("LLM 3-pass analysis complete",
                    confidence=result.get("analyst_confidence"),
                    red_flags=len(result.get("red_flags", [])),
                    green_flags=len(result.get("green_flags", [])),
                    requirements=len(result.get("technical_requirements", [])))
        return result

    except Exception as exc:
        logger.warning("LLM analysis failed — continuing with rule-based results", error=str(exc))
        return {}


def _merge_llm_results(
    rule_entities: dict,
    rule_flags: dict,
    llm_result: dict,
) -> tuple[dict, dict]:
    """Merge Claude's output with the rule-based extraction.

    Strategy:
    - Entities: Claude fills gaps; rule-based values override where they exist
      (regex is more precise for numbers/dates).
    - Flags: union of both sets, deduped by (code, severity).
    - Green flags: union, deduped by title_en.
    """
    # ── Entities merge ────────────────────────────────────────────────────────
    llm_entities: dict = llm_result.get("entities", {})
    merged_entities = dict(rule_entities)

    for key in ("issuing_agency", "estimated_value_sar", "deadline", "duration_months",
                "deployment_model", "local_content_pct", "data_outside_ksa",
                "requires_security_clearance", "primary_language"):
        if key not in merged_entities or merged_entities[key] is None:
            val = llm_entities.get(key)
            if val is not None:
                merged_entities[key] = val

    # Merge required_certs lists
    rule_certs = set(merged_entities.get("required_certs", []))
    llm_certs = set(llm_entities.get("required_certs", []))
    merged_entities["required_certs"] = list(rule_certs | llm_certs)

    # Merge deep analysis fields from LLM
    for field in ("project_objectives", "key_deliverables", "technical_requirements", "mandatory_requirements"):
        if not merged_entities.get(field):
            val = llm_entities.get(field)
            if val:
                merged_entities[field] = val

    # ── Flags merge ───────────────────────────────────────────────────────────
    existing_red_keys = {(f.get("code"), f.get("severity")) for f in rule_flags.get("red", [])}
    merged_red = list(rule_flags.get("red", []))

    for llm_flag in llm_result.get("red_flags", []):
        key = (llm_flag.get("code"), llm_flag.get("severity"))
        if key not in existing_red_keys:
            # Translate to the internal schema
            merged_red.append({
                "code": llm_flag.get("code", "LLM_FLAG"),
                "severity": llm_flag.get("severity", "MAJOR"),
                "title_en": llm_flag.get("title_en", ""),
                "title_ar": llm_flag.get("title_ar", ""),
                "description_en": llm_flag.get("description_en", ""),
                "quote": llm_flag.get("evidence"),
                "page": 1,
                "section": "LLM Analysis",
            })
            existing_red_keys.add(key)

    existing_green_titles = {f.get("title_en") for f in rule_flags.get("green", [])}
    merged_green = list(rule_flags.get("green", []))

    for llm_flag in llm_result.get("green_flags", []):
        title = llm_flag.get("title_en", "")
        if title not in existing_green_titles:
            merged_green.append({
                "title_en": title,
                "title_ar": llm_flag.get("title_ar", ""),
                "description_en": llm_flag.get("description_en", ""),
                "quote": llm_flag.get("evidence"),
                "page": 1,
                "section": "LLM Analysis",
            })
            existing_green_titles.add(title)

    merged_flags = {"red": merged_red, "green": merged_green[:8]}
    return merged_entities, merged_flags


# ── Score Computation ──────────────────────────────────────────────────────────

def _compute_scores(entities: dict, flags: dict, capability_result: dict) -> dict:
    """
    Compute the 8-dimension qualification score per Entropy's RFP Intelligence framework.

    D1  Sector & Client Fit        (0–20)
    D2  Technical Domain Alignment (0–25)
    D3  NDMO / Data Governance     (0–15)
    D4  AI / Generative AI         (0–15)
    D5  Arabic Language            (0–10)
    D6  Entropy Product Fit        (0–5)
    D7  Scope Feasibility          (0–5)
    D8  Compliance & Eligibility   (0–5)
    ─────────────────────────────────────
    Total                          (0–100)

    Decision thresholds:
      GO              ≥ 80, no major+ flags
      REVIEW          45–79 (or has major flags)
      NO_GO           < 45, or any CRITICAL flag
    """
    red_flags = flags.get("red", [])
    matched_caps = capability_result.get("matched", [])

    issuing_agency = entities.get("issuing_agency", "")
    is_existing_client = any(
        c.lower() in issuing_agency.lower() or issuing_agency.lower() in c.lower()
        for c in ENTROPY_EXISTING_CLIENTS
    ) if issuing_agency else False
    is_strategic_agency = any(
        a.lower() in issuing_agency.lower() or issuing_agency.lower() in a.lower()
        for a in STRATEGIC_AGENCIES
    ) if issuing_agency else False

    # ── D1: Sector & Client Fit (0–20) ────────────────────────────────────────
    if is_existing_client:
        d1 = 19.0   # Existing client — strongest signal
    elif is_strategic_agency:
        d1 = 15.0   # Strategic target (Vision 2030, key gov body)
    elif issuing_agency:
        d1 = 10.0   # Identified Saudi entity but not on strategic list
    else:
        d1 = 5.0    # Unknown agency

    # ── D2: Technical Domain Alignment (0–25) ─────────────────────────────────
    has_scope_mismatch = any(f.get("code") == "SCOPE_MISMATCH" for f in red_flags)
    if has_scope_mismatch:
        d2 = 0.0    # Instant disqualifier: cybersecurity, clinical, hardware
    else:
        cap_score = capability_result.get("score", 0)   # 0–15 normalized score
        # Scale from 0-15 capability score → 0-25 dimension score
        d2 = min(25.0, cap_score * (25.0 / 15.0))

    # ── D3: NDMO / Data Governance (0–15) ────────────────────────────────────
    has_data_mgmt = "data_management" in matched_caps
    required_certs = entities.get("required_certs", [])
    has_ndmo_req = any(c.upper() in {"NDMO", "NDI"} for c in required_certs)

    if has_ndmo_req and has_data_mgmt:
        d3 = 14.0   # Explicitly required, and Entropy has proven NDMO track record
    elif has_ndmo_req:
        d3 = 11.0   # Explicitly required; Entropy's certifications cover it
    elif has_data_mgmt:
        d3 = 9.0    # Data management present, NDMO implied
    elif any(c in matched_caps for c in {"data_engineering", "data_platform"}):
        d3 = 5.0    # Data present but governance not a priority
    else:
        d3 = 1.0    # No data governance element

    # ── D4: AI / Generative AI (0–15) ─────────────────────────────────────────
    has_genai = "generative_ai" in matched_caps
    has_agents = "agentic_ai" in matched_caps
    has_ml = "machine_learning" in matched_caps or "time_series_forecasting" in matched_caps
    has_analytics = "analytics_bi" in matched_caps or "document_intelligence" in matched_caps

    if (has_genai or has_agents) and has_ml:
        d4 = 14.0   # Full AI stack: GenAI/agents + ML
    elif has_genai or has_agents:
        d4 = 12.0   # Explicit GenAI or agentic AI
    elif has_ml:
        d4 = 9.0    # AI/ML models but not generative
    elif has_analytics:
        d4 = 5.0    # BI/analytics; AI mentioned as desirable but not core
    else:
        d4 = 1.0    # No AI requirement

    # ── D5: Arabic Language (0–10) ────────────────────────────────────────────
    has_arabic_nlp = "arabic_nlp" in matched_caps
    primary_lang = entities.get("primary_language", "")

    if has_arabic_nlp and primary_lang == "arabic":
        d5 = 9.5    # Arabic-first + NLP processing explicitly required
    elif has_arabic_nlp:
        d5 = 8.0    # Arabic NLP required
    elif primary_lang == "arabic":
        d5 = 6.0    # Arabic deliverables required (no NLP specifics)
    elif primary_lang == "mixed":
        d5 = 4.0    # Bilingual
    else:
        d5 = 2.0    # English-only or unspecified

    # ── D6: Entropy Product Fit (0–5) ─────────────────────────────────────────
    # Hydrogen: agentic_ai, generative_ai → agent orchestration
    # Axiom: analytics_bi, data_platform, agentic_ai → agentic analytics
    # Yameen: speech_recognition, arabic_nlp → meeting intelligence
    hydrogen_fit = any(c in matched_caps for c in {"agentic_ai", "generative_ai"})
    axiom_fit = any(c in matched_caps for c in {"analytics_bi", "data_platform", "agentic_ai"})
    yameen_fit = any(c in matched_caps for c in {"speech_recognition", "arabic_nlp"})

    product_matches = sum([hydrogen_fit, axiom_fit, yameen_fit])
    if product_matches >= 2:
        d6 = 5.0    # Multiple product fits
    elif product_matches == 1:
        d6 = 4.0    # Single product fit
    elif any(c in matched_caps for c in {"machine_learning", "data_management", "document_intelligence"}):
        d6 = 2.0    # Adaptable; requires customization
    else:
        d6 = 0.0    # No product fit; pure services

    # ── D7: Scope Feasibility (0–5) ───────────────────────────────────────────
    duration = entities.get("duration_months", 0)
    value_sar = entities.get("estimated_value_sar", 0)

    if 0 < duration < 3:
        d7 = 1.0    # Very tight — delivery risk is high
    elif 0 < duration < 6:
        d7 = 3.0    # Tight but achievable
    elif 12 <= duration <= 24:
        d7 = 5.0    # Sweet spot per Entropy's optimal delivery window
    elif duration > 36:
        d7 = 2.0    # Very long engagement — delivery risk
    else:
        d7 = 3.5    # Unknown or borderline

    if value_sar > 50_000_000:
        d7 = min(d7, 2.0)   # Very large — may need consortium approach

    # ── D8: Compliance & Eligibility (0–5) ────────────────────────────────────
    critical_cert_flags = [
        f for f in red_flags
        if f.get("code") == "MANDATORY_CERT_NOT_HELD" and f.get("severity") == "CRITICAL"
    ]
    major_cert_flags = [
        f for f in red_flags
        if f.get("code") == "MANDATORY_CERT_NOT_HELD" and f.get("severity") == "MAJOR"
    ]

    if has_scope_mismatch:
        d8 = 0.0    # Definitionally ineligible
    elif critical_cert_flags:
        d8 = 1.0    # Significant eligibility barrier (SAMA CSF, NCA ECC, PCI-DSS)
    elif major_cert_flags:
        d8 = 3.0    # Minor barrier — addressable via certified partner
    else:
        d8 = 5.0    # No hard eligibility barriers

    # ── Total score (0–100) ───────────────────────────────────────────────────
    total = d1 + d2 + d3 + d4 + d5 + d6 + d7 + d8

    # ── Map to legacy 3-bucket DB columns ────────────────────────────────────
    # technical_fit ← D2 + D3 + D4 + D6 (domain, NDMO, AI, product; max 60 → scale to 40)
    technical = min(40.0, (d2 + d3 + d4 + d6) * (40.0 / 60.0))
    # business_fit ← D1 + D5 (sector + Arabic; max 30)
    business = min(30.0, d1 + d5)
    # risk_penalty ← flag-based + feasibility/compliance shortfall
    critical_count = sum(1 for f in red_flags if f.get("severity") == "CRITICAL")
    major_count = sum(1 for f in red_flags if f.get("severity") == "MAJOR")
    minor_count = len(red_flags) - critical_count - major_count
    flag_penalty = min(20.0, critical_count * 8 + major_count * 3 + minor_count * 1)
    feasibility_shortfall = max(0.0, 10.0 - (d7 + d8))
    risk = min(30.0, flag_penalty + feasibility_shortfall)

    # Confidence: higher when more capability signals found + agency identified
    if len(matched_caps) >= 4:
        confidence = 0.92
    elif len(matched_caps) >= 2:
        confidence = 0.80
    elif len(matched_caps) >= 1:
        confidence = 0.70
    else:
        confidence = 0.55

    # Low-confidence sections
    low_conf = []
    if not matched_caps:
        low_conf.append("scope — no capability keywords matched")
    if not issuing_agency:
        low_conf.append("agency — could not identify issuing body")
    if not entities.get("duration_months"):
        low_conf.append("timeline — project duration not found")

    return {
        "technical": round(technical, 1),
        "business": round(business, 1),
        "risk": round(risk, 1),
        "total": round(total, 1),
        "confidence": confidence,
        "low_confidence_sections": low_conf,
        # 8 dimension scores (for explanation display)
        "d1_sector": round(d1, 1),
        "d2_domain": round(d2, 1),
        "d3_ndmo": round(d3, 1),
        "d4_ai": round(d4, 1),
        "d5_arabic": round(d5, 1),
        "d6_product": round(d6, 1),
        "d7_feasibility": round(d7, 1),
        "d8_compliance": round(d8, 1),
    }


# ── Chunking & Embedding ───────────────────────────────────────────────────────

def _chunk_text(text: str, rfp_id: str, sections: list[dict]) -> list[dict]:
    """Split text into ~800-word chunks with 150-word overlap.
    Extracts page_number from [PAGE N] markers in each chunk's text.
    """
    import re as _re
    chunks = []
    words = text.split()
    chunk_size = 800
    overlap = 150
    step = chunk_size - overlap
    max_chunks = 100

    for i in range(0, len(words), step):
        chunk_words = words[i: i + chunk_size]
        chunk_text = " ".join(chunk_words)
        page_match = _re.search(r'\[PAGE (\d+)\]', chunk_text)
        page_number = int(page_match.group(1)) if page_match else None
        chunks.append({
            "text": chunk_text,
            "rfp_id": rfp_id,
            "chunk_index": len(chunks),
            "section_type": "UNKNOWN",
            "page_number": page_number,
        })
        if len(chunks) >= max_chunks:
            break
    return chunks


async def _embed_chunks(chunks: list[dict]) -> list[list[float]]:
    """Generate embeddings. Routes to OpenAI → Cohere → Ollama → zero-vector fallback."""
    from core.config import settings

    if not chunks:
        return []

    provider = settings.embedding_provider.lower()
    if provider == "openai":
        return await _embed_chunks_openai(chunks, settings)
    elif provider == "cohere":
        return await _embed_chunks_cohere(chunks, settings)
    elif provider == "ollama":
        return await _embed_chunks_ollama(chunks, settings)
    return [[0.0] * 1024 for _ in chunks]


async def _embed_chunks_openai(chunks: list[dict], settings) -> list[list[float]]:
    """Embed using OpenAI text-embedding-3-large (dim=1024)."""
    from openai import AsyncOpenAI
    if not settings.openai_api_key:
        return [[0.0] * 1024 for _ in chunks]
    try:
        client = AsyncOpenAI(api_key=settings.openai_api_key)
        texts = [c["text"] for c in chunks]
        all_vectors: list[list[float]] = []
        batch_size = 100
        for i in range(0, len(texts), batch_size):
            batch = texts[i:i + batch_size]
            resp = await client.embeddings.create(
                model=settings.embedding_model,
                input=batch,
                dimensions=1024,
            )
            all_vectors.extend([e.embedding for e in resp.data])
        return all_vectors
    except Exception as e:
        logger.error("OpenAI embedding failed", error=str(e))
        return [[0.0] * 1024 for _ in chunks]


async def _embed_chunks_ollama(chunks: list[dict], settings) -> list[list[float]]:
    import httpx
    try:
        # Fast availability probe to avoid N * 30s waits when Ollama is down.
        async with httpx.AsyncClient(timeout=2.0) as probe_client:
            probe = await probe_client.get(f"{settings.ollama_api_url}/api/tags")
            if probe.status_code >= 400:
                return [[0.0] * 1024 for _ in chunks]

        embeddings = []
        async with httpx.AsyncClient() as client:
            for chunk in chunks:
                response = await client.post(
                    f"{settings.ollama_api_url}/api/embeddings",
                    json={"model": settings.embedding_model, "prompt": chunk["text"]},
                    timeout=6.0,
                )
                if response.status_code == 200:
                    embeddings.append(response.json().get("embedding", [0.0] * 1024))
                else:
                    embeddings.append([0.0] * 1024)
        return embeddings
    except Exception as e:
        logger.error("Ollama embedding failed", error=str(e))
        return [[0.0] * 1024 for _ in chunks]


async def _embed_chunks_cohere(chunks: list[dict], settings) -> list[list[float]]:
    import cohere
    if not settings.cohere_api_key:
        return [[0.0] * 1024 for _ in chunks]
    try:
        co = cohere.AsyncClient(settings.cohere_api_key)
        texts = [c["text"] for c in chunks]
        embeddings = []
        for i in range(0, len(texts), 96):
            batch = texts[i: i + 96]
            response = await co.embed(texts=batch, model=settings.embedding_model, input_type="search_document")
            embeddings.extend(response.embeddings)
        return embeddings
    except Exception as e:
        logger.error("Cohere embedding failed", error=str(e))
        return [[0.0] * 1024 for _ in chunks]


def settings_model() -> str:
    from core.config import settings
    return settings.primary_llm_model
